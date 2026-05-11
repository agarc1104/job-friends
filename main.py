import flet_webview
import flet as ft
import requests
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import asyncio
import base64
import hashlib
import json
import os
import re
import traceback
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from supabase import create_client, Client, ClientOptions
import httpx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas
from backend_config import settings
from services.job_search_service import search_google_jobs

if not settings.supabase_url or not settings.supabase_key:
    raise RuntimeError("Configura SUPABASE_URL y SUPABASE_KEY en el entorno antes de iniciar la app.")

SUPABASE_URL = settings.supabase_url
SUPABASE_KEY = settings.supabase_key

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# auxiliar para hashear contraseñas

def hash_password(plain: str) -> str:
    return hashlib.sha256(plain.encode("utf-8")).hexdigest()

# detalles en memoria para personalizar saludo
user_details: dict[str, dict[str, str]] = {}
# usuario logueado (email)
current_user: str | None = None
interview_chat_state_by_application: dict[str, list[dict[str, str]]] = {}
pending_interview_application_id: str | None = None

# archivo local para recordar credenciales
CREDENTIALS_FILE = "saved_credentials.json"
CV_BUCKET = "applicant-cvs"
GEMINI_API_KEY = settings.gemini_api_key
GEMINI_MODEL = settings.gemini_model
GEMINI_API_VERSION = settings.gemini_api_version
GEMINI_FALLBACK_MODELS = settings.gemini_fallback_models
CV_METADATA_TABLE = settings.cv_metadata_table
ALLOWED_CV_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
AI_CV_OUTPUT_FORMATS = {"docx", "pdf"}
ALLOWED_PHOTO_EXTENSIONS = {".png", ".jpg", ".jpeg"}
MAX_PROFILE_PHOTO_BYTES = 3 * 1024 * 1024
AI_CV_COLOR_PALETTES: list[dict[str, str]] = [
    {
        "id": "azul_profesional",
        "name": "Azul profesional",
        "description": "Paleta sobria azul-gris para un look corporativo y confiable.",
    },
    {
        "id": "verde_moderno",
        "name": "Verde moderno",
        "description": "Estilo limpio y contemporáneo con acentos verdes para destacar secciones.",
    },
    {
        "id": "gris_ejecutivo",
        "name": "Gris ejecutivo",
        "description": "Diseño elegante y neutral enfocado en una lectura formal.",
    },
]

AI_CV_FONT_SIZE_OPTIONS: list[dict[str, str]] = [
    {
        "id": "compacta",
        "name": "Compacta",
        "description": "Texto ligeramente más pequeño para aprovechar mejor el espacio.",
    },
    {
        "id": "estandar",
        "name": "Estandar",
        "description": "Tamaño equilibrado recomendado para la mayoría de vacantes.",
    },
    {
        "id": "amplia",
        "name": "Amplia",
        "description": "Texto más grande para máxima legibilidad.",
    },
]

AI_CV_COLUMN_OPTIONS: list[dict[str, str]] = [
    {
        "id": "una_columna",
        "name": "1 columna",
        "description": "Flujo lineal ATS-friendly con lectura simple.",
    },
    {
        "id": "dos_columnas",
        "name": "2 columnas",
        "description": "Distribución moderna para separar experiencia y habilidades.",
    },
]

AI_CV_PHOTO_OPTIONS: list[dict[str, str]] = [
    {
        "id": "con_foto",
        "name": "Incluir foto",
        "description": "Reservar espacio para foto de perfil profesional.",
    },
    {
        "id": "sin_foto",
        "name": "Sin foto",
        "description": "CV sin fotografía para enfoque completamente textual.",
    },
]


def normalize_ai_cv_visual_choice(
    raw_value: str | None,
    options: list[dict[str, str]],
    default_id: str,
) -> str:
    normalized = normalize_text(raw_value or "").lower()
    valid_ids = {item["id"] for item in options}
    return normalized if normalized in valid_ids else default_id


def get_ai_cv_visual_label(selected_id: str | None, options: list[dict[str, str]], default_id: str) -> str:
    resolved_id = normalize_ai_cv_visual_choice(selected_id, options, default_id)
    for item in options:
        if item["id"] == resolved_id:
            return item["name"]
    return options[0]["name"] if options else ""


def normalize_ai_cv_color_palette(color_palette: str | None) -> str:
    return normalize_ai_cv_visual_choice(color_palette, AI_CV_COLOR_PALETTES, "azul_profesional")


def normalize_ai_cv_font_size(font_size: str | None) -> str:
    return normalize_ai_cv_visual_choice(font_size, AI_CV_FONT_SIZE_OPTIONS, "estandar")


def normalize_ai_cv_columns(columns: str | None) -> str:
    return normalize_ai_cv_visual_choice(columns, AI_CV_COLUMN_OPTIONS, "una_columna")


def normalize_ai_cv_photo_option(photo_option: str | None) -> str:
    return normalize_ai_cv_visual_choice(photo_option, AI_CV_PHOTO_OPTIONS, "sin_foto")


def get_ai_cv_color_palette_label(color_palette: str | None) -> str:
    return get_ai_cv_visual_label(color_palette, AI_CV_COLOR_PALETTES, "azul_profesional")


def get_ai_cv_font_size_label(font_size: str | None) -> str:
    return get_ai_cv_visual_label(font_size, AI_CV_FONT_SIZE_OPTIONS, "estandar")


def get_ai_cv_columns_label(columns: str | None) -> str:
    return get_ai_cv_visual_label(columns, AI_CV_COLUMN_OPTIONS, "una_columna")


def get_ai_cv_photo_option_label(photo_option: str | None) -> str:
    return get_ai_cv_visual_label(photo_option, AI_CV_PHOTO_OPTIONS, "sin_foto")


def save_credentials(email: str, password: str) -> None:
    try:
        with open(CREDENTIALS_FILE, "w") as f:
            json.dump({"email": email, "password": password}, f)
    except Exception:
        pass


def load_credentials() -> dict | None:
    if os.path.exists(CREDENTIALS_FILE):
        try:
            with open(CREDENTIALS_FILE) as f:
                return json.load(f)
        except Exception:
            return None
    return None


def remove_saved_credentials() -> None:
    try:
        if os.path.exists(CREDENTIALS_FILE):
            os.remove(CREDENTIALS_FILE)
    except Exception:
        pass


def sanitize_storage_name(value: str) -> str:
    normalized = normalize_text(value).lower() if value else ""
    normalized = re.sub(r"[^a-z0-9._-]+", "-", normalized)
    normalized = normalized.strip("-._")
    return normalized or "file"


def get_cv_content_type(file_name: str) -> str:
    suffix = Path(file_name or "").suffix.lower()
    return ALLOWED_CV_EXTENSIONS.get(suffix, "application/octet-stream")


def is_supported_cv_file(file_name: str) -> bool:
    return Path(file_name or "").suffix.lower() in ALLOWED_CV_EXTENSIONS


def is_supported_photo_file(file_name: str) -> bool:
    return Path(file_name or "").suffix.lower() in ALLOWED_PHOTO_EXTENSIONS


def extract_profile_photo_bytes(profile_data: dict[str, str]) -> bytes | None:
    raw_base64 = normalize_text(profile_data.get("cv_photo_base64", ""))
    if not raw_base64:
        return None
    try:
        return base64.b64decode(raw_base64)
    except Exception:
        return None


def normalize_ai_cv_output_format(output_format: str | None) -> str:
    normalized = normalize_text(output_format or "").lower()
    return normalized if normalized in AI_CV_OUTPUT_FORMATS else "docx"


def build_gemini_cv_prompt(profile_data: dict[str, str], full_name: str = "", email: str = "") -> str:
    target_roles = profile_data.get("target_roles", "")
    summary = profile_data.get("summary", "")
    skills = profile_data.get("skills", "")
    experience = profile_data.get("experience", "")
    education = profile_data.get("education", "")
    languages = profile_data.get("languages", "")
    certifications = profile_data.get("certifications", "")
    achievements = profile_data.get("achievements", "")
    selected_palette = normalize_ai_cv_color_palette(profile_data.get("cv_color_palette", ""))
    selected_font_size = normalize_ai_cv_font_size(profile_data.get("cv_font_size", ""))
    selected_columns = normalize_ai_cv_columns(profile_data.get("cv_columns", ""))
    selected_photo = normalize_ai_cv_photo_option(profile_data.get("cv_include_photo", ""))
    generation_mode = profile_data.get("cv_generation_mode", "")
    job_title = profile_data.get("job_title", "")
    job_company = profile_data.get("job_company", "")
    job_location = profile_data.get("job_location", "")
    job_description = profile_data.get("job_description", "")
    existing_cv_file_name = profile_data.get("existing_cv_file_name", "")
    existing_cv_public_url = profile_data.get("existing_cv_public_url", "")
    existing_cv_source = profile_data.get("existing_cv_source", "")

    generation_mode_label = {
        "adapt_existing": "Adaptar CV guardado y reemplazarlo",
        "new_from_form": "Generar CV nuevo desde formulario",
    }.get(generation_mode, generation_mode or "Generacion estandar")

    return (
        "Actua como un arquitecto de documentos para RRHH. "
        "Genera el CODIGO ESTRUCTURADO del CV en formato JSON valido para luego renderizar DOCX o PDF.\n\n"
        f"Nombre completo: {full_name}\n"
        f"Email: {email}\n"
        f"Modo de generacion: {generation_mode_label}\n"
        f"Vacantes objetivo: {target_roles}\n"
        f"Resumen profesional: {summary}\n"
        f"Habilidades: {skills}\n"
        f"Experiencia: {experience}\n"
        f"Educacion: {education}\n"
        f"Idiomas: {languages}\n"
        f"Certificaciones: {certifications}\n"
        f"Logros: {achievements}\n"
        f"Vacante actual: {job_title} en {job_company}\n"
        f"Ubicacion de la vacante: {job_location}\n"
        f"Descripcion completa de la vacante: {job_description}\n"
        f"CV guardado en sistema: {existing_cv_file_name}\n"
        f"Origen del CV guardado: {existing_cv_source}\n"
        f"Referencia del CV guardado: {existing_cv_public_url}\n"
        f"Paleta de color: {get_ai_cv_color_palette_label(selected_palette)} ({selected_palette})\n"
        f"Tamano de letra: {get_ai_cv_font_size_label(selected_font_size)} ({selected_font_size})\n"
        f"Distribucion: {get_ai_cv_columns_label(selected_columns)} ({selected_columns})\n"
        f"Foto en CV: {get_ai_cv_photo_option_label(selected_photo)} ({selected_photo})\n\n"
        "Instrucciones de adaptacion:\n"
        "- Usa la descripcion completa de la vacante para priorizar palabras clave, habilidades y logros relevantes.\n"
        "- Si el modo es adaptar CV guardado, conserva la trayectoria del candidato y optimizala para la vacante sin inventar experiencia.\n"
        "- Si el modo es generar CV nuevo desde formulario, usa primero la informacion escrita en el formulario y complementala con el enfoque de la vacante.\n"
        "- Si existe referencia a un CV guardado, tomalo como contexto adicional del perfil ya registrado, pero produce una nueva version completa para reemplazo.\n\n"
        "Reglas de salida obligatorias:\n"
        "1) Devuelve SOLO JSON valido, sin markdown, sin explicaciones y sin texto fuera del JSON.\n"
        "2) Usa exactamente esta estructura de alto nivel:\n"
        "{\n"
        "  \"document_type\": \"cv_blueprint_v1\",\n"
        "  \"header\": {\"full_name\": \"...\", \"email\": \"...\"},\n"
        "  \"sections\": [\n"
        "    {\"id\": \"target_roles\", \"title\": \"Vacantes objetivo\", \"items\": [\"...\"]},\n"
        "    {\"id\": \"profile\", \"title\": \"Perfil profesional\", \"items\": [\"...\"]},\n"
        "    {\"id\": \"experience\", \"title\": \"Experiencia\", \"items\": [\"...\"]},\n"
        "    {\"id\": \"education\", \"title\": \"Educacion\", \"items\": [\"...\"]},\n"
        "    {\"id\": \"skills\", \"title\": \"Habilidades\", \"items\": [\"...\"]},\n"
        "    {\"id\": \"languages\", \"title\": \"Idiomas\", \"items\": [\"...\"]},\n"
        "    {\"id\": \"certifications\", \"title\": \"Certificaciones\", \"items\": [\"...\"]},\n"
        "    {\"id\": \"achievements\", \"title\": \"Logros\", \"items\": [\"...\"]}\n"
        "  ],\n"
        "  \"render_hints\": {\"target_formats\": [\"docx\", \"pdf\"], \"language\": \"es\"}\n"
        "}\n"
        "3) Incluye SIEMPRE las 8 secciones anteriores, aunque alguna tenga 1 item breve.\n"
        "4) Manten tono profesional, concreto y ATS-friendly.\n"
        "5) Ajusta la redaccion y jerarquia visual a las preferencias de diseno sin cambiar la estructura JSON obligatoria."
    )


CV_SECTION_CANONICAL = {
    "vacantes objetivo": "Vacantes objetivo",
    "perfil": "Perfil profesional",
    "perfil profesional": "Perfil profesional",
    "resumen profesional": "Perfil profesional",
    "experiencia": "Experiencia",
    "educacion": "Educacion",
    "habilidades": "Habilidades",
    "idiomas": "Idiomas",
    "certificaciones": "Certificaciones",
    "logros": "Logros",
}

CV_SECTION_ORDER: list[tuple[str, str, str]] = [
    ("target_roles", "Vacantes objetivo", "target_roles"),
    ("profile", "Perfil profesional", "summary"),
    ("experience", "Experiencia", "experience"),
    ("education", "Educacion", "education"),
    ("skills", "Habilidades", "skills"),
    ("languages", "Idiomas", "languages"),
    ("certifications", "Certificaciones", "certifications"),
    ("achievements", "Logros", "achievements"),
]


def canonicalize_cv_heading(line: str) -> str | None:
    normalized = normalize_text(line).lower().strip().rstrip(":")
    return CV_SECTION_CANONICAL.get(normalized)


def normalize_cv_items(raw_items: str | list | None) -> list[str]:
    if isinstance(raw_items, str):
        candidate_items = [segment.strip() for segment in re.split(r"\n|;", raw_items) if segment.strip()]
        return [normalize_text(item) for item in candidate_items if normalize_text(item)]
    if isinstance(raw_items, list):
        output: list[str] = []
        for item in raw_items:
            normalized = normalize_text(str(item)) if item is not None else ""
            if normalized:
                output.append(normalized)
        return output
    return []


def parse_gemini_json_output(raw_text: str) -> dict:
    text = (raw_text or "").replace("\r", "").strip()
    if not text:
        raise ValueError("Gemini devolvio salida vacia para el CV estructurado.")

    # Quitar fences markdown en caso de respuestas con ```
    text = re.sub(r"^```[a-zA-Z0-9_-]*\n", "", text)
    text = re.sub(r"\n```$", "", text)

    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    first = text.find("{")
    last = text.rfind("}")
    if first != -1 and last != -1 and first < last:
        candidate = text[first : last + 1]
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass

    raise ValueError("No se pudo parsear JSON valido en la salida de Gemini.")


def build_local_cv_structure(profile_data: dict[str, str], full_name: str, email: str) -> dict:
    sections: list[dict] = []
    for section_id, title, profile_key in CV_SECTION_ORDER:
        content = normalize_text(profile_data.get(profile_key, ""))
        sections.append(
            {
                "id": section_id,
                "title": title,
                "items": normalize_cv_items(content) or ["Sin informacion proporcionada."],
            }
        )

    return {
        "document_type": "cv_blueprint_v1",
        "header": {
            "full_name": normalize_text(full_name) or "Curriculum Vitae",
            "email": normalize_text(email),
        },
        "sections": sections,
        "render_hints": {"target_formats": ["docx", "pdf"], "language": "es"},
    }


def normalize_cv_structure(raw_cv: dict, profile_data: dict[str, str], full_name: str, email: str) -> dict:
    fallback = build_local_cv_structure(profile_data, full_name, email)
    if not isinstance(raw_cv, dict):
        return fallback

    header = raw_cv.get("header") if isinstance(raw_cv.get("header"), dict) else {}
    merged_header = {
        "full_name": normalize_text(str(header.get("full_name", ""))) or fallback["header"]["full_name"],
        "email": normalize_text(str(header.get("email", ""))) or fallback["header"]["email"],
    }

    existing_sections = raw_cv.get("sections") if isinstance(raw_cv.get("sections"), list) else []
    by_id: dict[str, dict] = {}
    by_title: dict[str, dict] = {}
    for section in existing_sections:
        if not isinstance(section, dict):
            continue
        sec_id = normalize_text(str(section.get("id", ""))).lower()
        sec_title = normalize_text(str(section.get("title", "")))
        if sec_id:
            by_id[sec_id] = section
        canonical_title = canonicalize_cv_heading(sec_title)
        if canonical_title:
            by_title[canonical_title] = section

    normalized_sections: list[dict] = []
    for section_id, title, profile_key in CV_SECTION_ORDER:
        source = by_id.get(section_id) or by_title.get(title)
        raw_items = source.get("items") if isinstance(source, dict) else None
        normalized_items = normalize_cv_items(raw_items)
        if not normalized_items:
            local_content = normalize_text(profile_data.get(profile_key, ""))
            normalized_items = normalize_cv_items(local_content) or ["Sin informacion proporcionada."]

        normalized_sections.append({"id": section_id, "title": title, "items": normalized_items})

    render_hints = raw_cv.get("render_hints") if isinstance(raw_cv.get("render_hints"), dict) else {}
    formats = render_hints.get("target_formats") if isinstance(render_hints.get("target_formats"), list) else ["docx", "pdf"]
    normalized_formats = [normalize_text(str(item)).lower() for item in formats if normalize_text(str(item))]
    if not normalized_formats:
        normalized_formats = ["docx", "pdf"]

    return {
        "document_type": "cv_blueprint_v1",
        "header": merged_header,
        "sections": normalized_sections,
        "render_hints": {
            "target_formats": normalized_formats,
            "language": normalize_text(str(render_hints.get("language", "es"))) or "es",
        },
    }


def looks_like_structured_cv_code(cv_data: dict) -> bool:
    if not isinstance(cv_data, dict):
        return False
    sections = cv_data.get("sections")
    if not isinstance(sections, list):
        return False
    existing_ids = {
        normalize_text(str(item.get("id", ""))).lower()
        for item in sections
        if isinstance(item, dict)
    }
    required_ids = {item[0] for item in CV_SECTION_ORDER}
    return required_ids.issubset(existing_ids)


def extract_cv_html_from_gemini_text(raw_text: str) -> str:
    """Extrae el bloque HTML de la respuesta de Gemini, eliminando markdown code fences si los hay."""
    for pattern in [r"```html\s*([\s\S]*?)\s*```", r"```\s*([\s\S]*?)\s*```"]:
        match = re.search(pattern, raw_text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return raw_text.strip()


def looks_like_cv_html(html: str) -> bool:
    """Valida que la cadena tenga la forma de un documento HTML de CV."""
    if not html or len(html) < 300:
        return False
    lower = html.lower()
    return "<html" in lower and ("<body" in lower or "<head" in lower) and (
        "<h1" in lower or "<h2" in lower or "<section" in lower or "<div" in lower
    )


def inject_photo_into_html(html: str, photo_bytes: bytes) -> str:
    """Inyecta la foto del usuario como data URI dentro del elemento img#cv-photo."""
    import mimetypes as _mt
    mime = "image/jpeg"
    b64 = base64.b64encode(photo_bytes).decode("utf-8")
    data_uri = f"data:{mime};base64,{b64}"
    if 'id="cv-photo"' in html:
        html = re.sub(
            r'(<img[^>]*id="cv-photo"[^>]*?)(?:\ssrc="[^"]*")?(\s*/?>)',
            lambda m: f'{m.group(1)} src="{data_uri}"{m.group(2)}',
            html,
            count=1,
        )
    return html


def _sanitize_html_for_xhtml2pdf(html: str) -> str:
    """Elimina o reemplaza construcciones CSS no soportadas por xhtml2pdf/pisa."""
    # Eliminar bloques ::before y ::after (ej. li::before { content: "•"; ... })
    html = re.sub(r"[^{}\s]*::(?:before|after)\s*\{[^}]*\}", "", html, flags=re.DOTALL | re.IGNORECASE)
    # Reemplazar display:flex por display:block (xhtml2pdf no entiende flexbox)
    html = re.sub(r"display\s*:\s*flex\s*;?", "display: block;", html, flags=re.IGNORECASE)
    # Eliminar propiedades exclusivas de flexbox
    for prop in (
        r"gap", r"flex(?:-direction|-wrap|-grow|-shrink|-basis)?",
        r"align-items", r"align-content", r"justify-content",
    ):
        html = re.sub(rf"\b{prop}\s*:[^;}}]+;?", "", html, flags=re.IGNORECASE)
    # Convertir .content-columns con columnas a tabla simple
    # (Gemini suele usar esta clase para layout de 2 columnas)
    html = re.sub(
        r"(\.content-columns\s*\{[^}]*?)display\s*:\s*block\s*;",
        r"\1",
        html,
        flags=re.DOTALL | re.IGNORECASE,
    )
    # Sustituir li sin viñeta por list-style disc como fallback legible
    html = re.sub(
        r"(ul\s*\{[^}]*?)list-style-type\s*:\s*none\s*;",
        r"\1list-style-type: disc;",
        html,
        flags=re.DOTALL | re.IGNORECASE,
    )
    return html


def create_cv_pdf_from_html(html: str) -> bytes:
    """Convierte una cadena HTML a bytes de PDF usando xhtml2pdf."""
    from xhtml2pdf import pisa
    sanitized = _sanitize_html_for_xhtml2pdf(html)
    buffer = BytesIO()
    html_bytes = sanitized.encode("utf-8")
    pisa.pisaDocument(BytesIO(html_bytes), buffer, encoding="utf-8")
    return buffer.getvalue()


def create_cv_docx_from_html(html: str) -> bytes:
    """Convierte una cadena HTML a bytes de DOCX usando htmldocx."""
    from htmldocx import HtmlToDocx
    document = Document()
    parser = HtmlToDocx()
    parser.add_html_to_document(html, document)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def extract_gemini_text(response_data: dict) -> str:
    candidates = response_data.get("candidates", []) or []
    text_parts: list[str] = []

    for candidate in candidates:
        content = candidate.get("content", {}) or {}
        parts = content.get("parts", []) or []
        for part in parts:
            text = part.get("text", "")
            if text:
                text_parts.append(text)

    return "\n".join(text_parts).strip()


def extract_gemini_error_detail(response: httpx.Response | None) -> str:
    if response is None:
        return ""
    try:
        payload = response.json()
    except ValueError:
        return ""
    error_obj = payload.get("error", {}) if isinstance(payload, dict) else {}
    message = error_obj.get("message", "") if isinstance(error_obj, dict) else ""
    return str(message or "").strip()


def is_supabase_missing_table_error(ex: Exception, table_name: str) -> bool:
    message = str(ex)
    normalized = message.lower()
    return "pgrst205" in normalized and table_name.lower() in normalized


def generate_cv_structure_with_gemini(profile_data: dict[str, str], full_name: str, email: str) -> dict:
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY no configurada.")

    prompt = build_gemini_cv_prompt(profile_data, full_name, email)
    payload = {
        "contents": [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0.5,
            "topP": 0.9,
            "maxOutputTokens": 2048,
        },
    }

    models_to_try = [GEMINI_MODEL] + [m for m in GEMINI_FALLBACK_MODELS if m != GEMINI_MODEL]
    last_error_detail = ""
    last_status_code = 0

    for model_name in models_to_try:
        url = f"https://generativelanguage.googleapis.com/{GEMINI_API_VERSION}/models/{model_name}:generateContent"
        try:
            response = httpx.post(
                url,
                json=payload,
                headers={"x-goog-api-key": GEMINI_API_KEY},
                timeout=45.0,
            )
            response.raise_for_status()
            generated_text = extract_gemini_text(response.json())
            if not generated_text:
                raise ValueError("Gemini no devolvio contenido para el CV estructurado.")

            parsed_cv = parse_gemini_json_output(generated_text)
            normalized_cv = normalize_cv_structure(parsed_cv, profile_data, full_name, email)

            if not looks_like_structured_cv_code(normalized_cv):
                repair_prompt = (
                    "Repara y completa el siguiente JSON de CV. "
                    "Responde SOLO con JSON valido, sin markdown, con secciones obligatorias: "
                    "target_roles, profile, experience, education, skills, languages, certifications, achievements.\n\n"
                    f"JSON a reparar:\n{generated_text}"
                )
                repair_payload = {
                    "contents": [
                        {
                            "parts": [
                                {"text": repair_prompt}
                            ]
                        }
                    ],
                    "generationConfig": {
                        "temperature": 0.3,
                        "topP": 0.8,
                        "maxOutputTokens": 2048,
                    },
                }
                repair_response = httpx.post(
                    url,
                    json=repair_payload,
                    headers={"x-goog-api-key": GEMINI_API_KEY},
                    timeout=45.0,
                )
                repair_response.raise_for_status()
                repaired_text = extract_gemini_text(repair_response.json())
                repaired_cv = parse_gemini_json_output(repaired_text)
                normalized_cv = normalize_cv_structure(repaired_cv, profile_data, full_name, email)

            if not looks_like_structured_cv_code(normalized_cv):
                raise ValueError("Gemini devolvio una estructura de CV incompleta tras normalizacion.")

            return normalized_cv
        except httpx.HTTPStatusError as ex:
            status_code = ex.response.status_code
            error_detail = extract_gemini_error_detail(ex.response)
            last_status_code = status_code
            last_error_detail = error_detail

            if status_code == 404:
                continue
            if status_code == 429:
                raise ValueError("Gemini devolvio 429 por limite de cuota o solicitudes.") from None
            if error_detail:
                raise ValueError(f"Gemini devolvio error HTTP {status_code}: {error_detail}") from None
            raise ValueError(f"Gemini devolvio error HTTP {status_code}.") from None
        except httpx.HTTPError as ex:
            raise ValueError(f"No fue posible conectar con Gemini: {ex}") from None

    attempted_models = ", ".join(models_to_try)
    if last_status_code == 404:
        detail_suffix = f" Detalle: {last_error_detail}" if last_error_detail else ""
        raise ValueError(
            "Gemini devolvio 404 para los modelos configurados. "
            f"Modelos probados: {attempted_models}.{detail_suffix}"
        )
    raise ValueError("No fue posible generar la estructura del CV con Gemini.")


def generate_cv_structure(profile_data: dict[str, str], full_name: str, email: str) -> dict:
    if GEMINI_API_KEY:
        try:
            return generate_cv_structure_with_gemini(profile_data, full_name, email)
        except Exception as ex:
            pass

    # Fallback local si Gemini no esta disponible.
    return build_local_cv_structure(profile_data, full_name, email)


def create_cv_docx_bytes(full_name: str, email: str, profile_data: dict[str, str], cv_structure: dict) -> bytes:
    def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
        cleaned = (hex_color or "000000").strip().lstrip("#")
        if len(cleaned) != 6:
            cleaned = "000000"
        return int(cleaned[0:2], 16), int(cleaned[2:4], 16), int(cleaned[4:6], 16)

    palette_map = {
        "azul_profesional": {"accent": "1F4E79", "text": "111827"},
        "verde_moderno": {"accent": "166534", "text": "0F172A"},
        "gris_ejecutivo": {"accent": "374151", "text": "111827"},
    }
    font_map = {"compacta": 10, "estandar": 11, "amplia": 12}

    selected_palette = normalize_ai_cv_color_palette(profile_data.get("cv_color_palette", ""))
    selected_font_size = normalize_ai_cv_font_size(profile_data.get("cv_font_size", ""))
    selected_columns = normalize_ai_cv_columns(profile_data.get("cv_columns", ""))
    include_photo = normalize_ai_cv_photo_option(profile_data.get("cv_include_photo", "")) == "con_foto"
    photo_bytes = extract_profile_photo_bytes(profile_data)

    colors = palette_map.get(selected_palette, palette_map["azul_profesional"])
    body_size = font_map.get(selected_font_size, 11)
    title_size = body_size + 8
    section_size = body_size + 2

    def _set_run_style(run, size: int, color_hex: str, bold: bool = False, italic: bool = False):
        r, g, b = _hex_to_rgb(color_hex)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(r, g, b)

    document = Document()
    normalized_structure = normalize_cv_structure(cv_structure, profile_data, full_name, email)

    header = normalized_structure.get("header", {}) if isinstance(normalized_structure, dict) else {}
    doc_name = normalize_text(str(header.get("full_name", ""))) or full_name or "Curriculum Vitae"
    doc_email = normalize_text(str(header.get("email", ""))) or email

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(doc_name)
    _set_run_style(title_run, size=title_size, color_hex=colors["accent"], bold=True)

    if doc_email:
        email_paragraph = document.add_paragraph()
        email_run = email_paragraph.add_run(doc_email)
        _set_run_style(email_run, size=body_size, color_hex=colors["text"])

    if include_photo:
        if photo_bytes:
            try:
                document.add_picture(BytesIO(photo_bytes), width=Inches(1.25))
            except Exception:
                photo_paragraph = document.add_paragraph()
                photo_run = photo_paragraph.add_run("[No se pudo cargar la foto]")
                _set_run_style(photo_run, size=body_size, color_hex=colors["accent"], italic=True)
        else:
            photo_paragraph = document.add_paragraph()
            photo_run = photo_paragraph.add_run("[Foto de perfil]")
            _set_run_style(photo_run, size=body_size, color_hex=colors["accent"], italic=True)

    sections = [section for section in normalized_structure.get("sections", []) or [] if isinstance(section, dict)]

    if selected_columns == "dos_columnas":
        split_index = (len(sections) + 1) // 2
        left_sections = sections[:split_index]
        right_sections = sections[split_index:]
        table = document.add_table(rows=1, cols=2)
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        left_cell.text = ""
        right_cell.text = ""

        def fill_cell(cell, grouped_sections):
            for section in grouped_sections:
                section_title = normalize_text(str(section.get("title", "")))
                if section_title:
                    title_p = cell.add_paragraph()
                    title_run = title_p.add_run(section_title)
                    _set_run_style(title_run, size=section_size, color_hex=colors["accent"], bold=True)

                for item in normalize_cv_items(section.get("items")):
                    item_p = cell.add_paragraph()
                    item_run = item_p.add_run(f"- {item}")
                    _set_run_style(item_run, size=body_size, color_hex=colors["text"])

        fill_cell(left_cell, left_sections)
        fill_cell(right_cell, right_sections)
        return_bytes = BytesIO()
        document.save(return_bytes)
        return return_bytes.getvalue()

    for section in sections:
        section_title = normalize_text(str(section.get("title", "")))
        if section_title:
            section_p = document.add_paragraph()
            section_run = section_p.add_run(section_title)
            _set_run_style(section_run, size=section_size, color_hex=colors["accent"], bold=True)
        for item in normalize_cv_items(section.get("items")):
            item_p = document.add_paragraph()
            item_run = item_p.add_run(f"- {item}")
            _set_run_style(item_run, size=body_size, color_hex=colors["text"])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def split_text_for_pdf(text: str, max_width: float, font_name: str, font_size: int) -> list[str]:
    normalized_text = normalize_text(text)
    if not normalized_text:
        return [""]

    words = normalized_text.split(" ")
    lines: list[str] = []
    current_line = ""

    for word in words:
        candidate = f"{current_line} {word}".strip()
        if candidate and pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
            current_line = candidate
            continue

        if current_line:
            lines.append(current_line)
        current_line = word

    if current_line:
        lines.append(current_line)

    return lines or [""]


def create_cv_pdf_bytes(full_name: str, email: str, profile_data: dict[str, str], cv_structure: dict) -> bytes:
    palette_map = {
        "azul_profesional": {"accent": (0.12, 0.31, 0.47), "text": (0.07, 0.09, 0.13)},
        "verde_moderno": {"accent": (0.09, 0.40, 0.20), "text": (0.06, 0.09, 0.16)},
        "gris_ejecutivo": {"accent": (0.22, 0.24, 0.27), "text": (0.07, 0.09, 0.13)},
    }
    font_map = {"compacta": 10, "estandar": 11, "amplia": 12}

    selected_palette = normalize_ai_cv_color_palette(profile_data.get("cv_color_palette", ""))
    selected_font_size = normalize_ai_cv_font_size(profile_data.get("cv_font_size", ""))
    selected_columns = normalize_ai_cv_columns(profile_data.get("cv_columns", ""))
    include_photo = normalize_ai_cv_photo_option(profile_data.get("cv_include_photo", "")) == "con_foto"
    photo_bytes = extract_profile_photo_bytes(profile_data)

    colors = palette_map.get(selected_palette, palette_map["azul_profesional"])
    body_size = font_map.get(selected_font_size, 11)
    section_size = body_size + 2
    title_size = body_size + 7

    normalized_structure = normalize_cv_structure(cv_structure, profile_data, full_name, email)
    header = normalized_structure.get("header", {}) if isinstance(normalized_structure, dict) else {}
    doc_name = normalize_text(str(header.get("full_name", ""))) or full_name or "Curriculum Vitae"
    doc_email = normalize_text(str(header.get("email", ""))) or email

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    page_width, page_height = LETTER
    margin = 50
    y = page_height - margin

    def ensure_space(required_height: int):
        nonlocal y
        if y - required_height < margin:
            pdf.showPage()
            y = page_height - margin

    def draw_wrapped_text(
        text: str,
        font_name: str,
        font_size: int,
        indent: int = 0,
        spacing: int = 15,
        color_rgb: tuple[float, float, float] = (0, 0, 0),
    ):
        nonlocal y
        max_width = page_width - (margin * 2) - indent
        lines = split_text_for_pdf(text, max_width=max_width, font_name=font_name, font_size=font_size)
        for line in lines:
            ensure_space(spacing)
            pdf.setFont(font_name, font_size)
            pdf.setFillColorRGB(*color_rgb)
            pdf.drawString(margin + indent, y, line)
            y -= spacing

    if include_photo:
        photo_w = 90
        photo_h = 110
        photo_x = page_width - margin - photo_w
        photo_y_top = page_height - margin
        if photo_bytes:
            try:
                image = ImageReader(BytesIO(photo_bytes))
                pdf.drawImage(
                    image,
                    photo_x,
                    photo_y_top - photo_h,
                    width=photo_w,
                    height=photo_h,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                pdf.setStrokeColorRGB(*colors["accent"])
                pdf.rect(photo_x, photo_y_top - photo_h, photo_w, photo_h, stroke=1, fill=0)
                pdf.setFillColorRGB(*colors["accent"])
                pdf.setFont("Helvetica", 10)
                pdf.drawCentredString(photo_x + (photo_w / 2), photo_y_top - (photo_h / 2), "Foto")
        else:
            pdf.setStrokeColorRGB(*colors["accent"])
            pdf.rect(photo_x, photo_y_top - photo_h, photo_w, photo_h, stroke=1, fill=0)
            pdf.setFillColorRGB(*colors["accent"])
            pdf.setFont("Helvetica", 10)
            pdf.drawCentredString(photo_x + (photo_w / 2), photo_y_top - (photo_h / 2), "Foto")

    draw_wrapped_text(
        doc_name,
        font_name="Helvetica-Bold",
        font_size=title_size,
        spacing=20,
        color_rgb=colors["accent"],
    )
    if doc_email:
        draw_wrapped_text(
            doc_email,
            font_name="Helvetica",
            font_size=body_size,
            spacing=16,
            color_rgb=colors["text"],
        )

    y -= 6

    sections = [section for section in normalized_structure.get("sections", []) or [] if isinstance(section, dict)]
    if selected_columns == "dos_columnas":
        column_gap = 22
        column_width = (page_width - (margin * 2) - column_gap) / 2
        column_x = [margin, margin + column_width + column_gap]
        column_y = [y, y]

        for idx, section in enumerate(sections):
            col = idx % 2
            x = column_x[col]
            y_col = column_y[col]

            section_title = normalize_text(str(section.get("title", "")))
            if section_title:
                title_lines = split_text_for_pdf(section_title, column_width, "Helvetica-Bold", section_size)
                for line in title_lines:
                    if y_col - 16 < margin:
                        pdf.showPage()
                        column_y = [page_height - margin, page_height - margin]
                        y_col = column_y[col]
                    pdf.setFont("Helvetica-Bold", section_size)
                    pdf.setFillColorRGB(*colors["accent"])
                    pdf.drawString(x, y_col, line)
                    y_col -= 16

            for item in normalize_cv_items(section.get("items")):
                item_lines = split_text_for_pdf(f"- {item}", column_width - 10, "Helvetica", body_size)
                for line in item_lines:
                    if y_col - 14 < margin:
                        pdf.showPage()
                        column_y = [page_height - margin, page_height - margin]
                        y_col = column_y[col]
                    pdf.setFont("Helvetica", body_size)
                    pdf.setFillColorRGB(*colors["text"])
                    pdf.drawString(x + 8, y_col, line)
                    y_col -= 14
            y_col -= 4
            column_y[col] = y_col
    else:
        for section in sections:
            section_title = normalize_text(str(section.get("title", "")))
            if section_title:
                ensure_space(22)
                draw_wrapped_text(
                    section_title,
                    font_name="Helvetica-Bold",
                    font_size=section_size,
                    spacing=17,
                    color_rgb=colors["accent"],
                )

            for item in normalize_cv_items(section.get("items")):
                draw_wrapped_text(
                    f"- {item}",
                    font_name="Helvetica",
                    font_size=body_size,
                    indent=12,
                    spacing=15,
                    color_rgb=colors["text"],
                )

            y -= 4

    pdf.save()
    return buffer.getvalue()


def get_supabase_public_url(storage_path: str) -> str:
    response = supabase.storage.from_(CV_BUCKET).get_public_url(storage_path)
    if isinstance(response, str):
        return response
    if isinstance(response, dict):
        return response.get("publicUrl") or response.get("public_url") or ""
    public_url = getattr(response, "public_url", None)
    if public_url:
        return public_url
    data = getattr(response, "data", None)
    if isinstance(data, dict):
        return data.get("publicUrl") or data.get("public_url") or ""
    return ""


def save_cv_metadata(
    email: str,
    file_name: str,
    storage_path: str,
    public_url: str,
    source: str,
    target_roles: str = "",
    profile_data: dict | None = None,
) -> dict:
    payload = {
        "applicant_email": email,
        "file_name": file_name,
        "storage_path": storage_path,
        "public_url": public_url,
        "source": source,
        "target_roles": target_roles,
        "profile_data": profile_data or {},
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }

    try:
        existing = supabase.table(CV_METADATA_TABLE).select("applicant_email").eq("applicant_email", email).execute()
        if existing.data:
            supabase.table(CV_METADATA_TABLE).update(payload).eq("applicant_email", email).execute()
        else:
            payload["created_at"] = payload["updated_at"]
            supabase.table(CV_METADATA_TABLE).insert(payload).execute()
    except Exception as ex:
        if is_supabase_missing_table_error(ex, CV_METADATA_TABLE):
            raise ValueError(
                f"No existe la tabla de CV '{CV_METADATA_TABLE}' en Supabase. "
                "Crea esa tabla o define SUPABASE_CV_TABLE con el nombre correcto."
            ) from None
        raise
    return payload


def normalize_profile_data(profile_data: dict | str | None) -> dict[str, str]:
    if isinstance(profile_data, dict):
        return {
            str(key): value if isinstance(value, str) else ("" if value is None else str(value))
            for key, value in profile_data.items()
        }

    if isinstance(profile_data, str):
        raw = profile_data.strip()
        if not raw:
            return {}
        try:
            decoded = json.loads(raw)
        except json.JSONDecodeError:
            return {}
        return normalize_profile_data(decoded)

    return {}


def upload_cv_bytes(
    email: str,
    file_name: str,
    content: bytes,
    source: str,
    target_roles: str = "",
    profile_data: dict | None = None,
) -> dict:
    if not content:
        raise ValueError("El archivo del CV no contiene datos.")

    safe_email = sanitize_storage_name(email)
    safe_file_name = sanitize_storage_name(Path(file_name).stem) + Path(file_name).suffix.lower()
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    storage_path = f"{safe_email}/{timestamp}_{safe_file_name}"
    content_type = get_cv_content_type(file_name)

    supabase.storage.from_(CV_BUCKET).upload(
        path=storage_path,
        file=content,
        file_options={"content-type": content_type},
    )

    public_url = get_supabase_public_url(storage_path)
    metadata = save_cv_metadata(
        email=email,
        file_name=file_name,
        storage_path=storage_path,
        public_url=public_url,
        source=source,
        target_roles=target_roles,
        profile_data=profile_data,
    )
    metadata["content_type"] = content_type
    return metadata


def upload_cv_file(email: str, file_path: str, file_name: str | None = None) -> dict:
    resolved_name = file_name or Path(file_path).name
    with open(file_path, "rb") as cv_file:
        content = cv_file.read()
    return upload_cv_bytes(email, resolved_name, content, source="uploaded")


def create_ai_cv_for_user(email: str, full_name: str, profile_data: dict[str, str], output_format: str = "docx") -> dict:
    normalized_output_format = normalize_ai_cv_output_format(output_format)
    enriched_profile = {
        key: value
        for key, value in profile_data.items()
        if key not in {"cv_photo_base64"}
    }
    enriched_profile["cv_photo_uploaded"] = bool(profile_data.get("cv_photo_base64"))
    enriched_profile["gemini_prompt"] = build_gemini_cv_prompt(profile_data, full_name, email)
    enriched_profile["gemini_ready"] = bool(GEMINI_API_KEY)
    enriched_profile["gemini_model"] = GEMINI_MODEL if GEMINI_API_KEY else ""
    enriched_profile["output_format"] = normalized_output_format
    cv_structure = generate_cv_structure(profile_data, full_name, email)
    cv_code = json.dumps(cv_structure, ensure_ascii=False, indent=2)
    if normalized_output_format == "pdf":
        file_bytes = create_cv_pdf_bytes(full_name, email, profile_data, cv_structure)
        file_extension = "pdf"
    else:
        file_bytes = create_cv_docx_bytes(full_name, email, profile_data, cv_structure)
        file_extension = "docx"
    file_name = f"cv_{sanitize_storage_name(full_name or email)}.{file_extension}"
    metadata = upload_cv_bytes(
        email=email,
        file_name=file_name,
        content=file_bytes,
        source="ai_generated",
        target_roles=profile_data.get("target_roles", ""),
        profile_data=enriched_profile,
    )
    metadata["output_format"] = normalized_output_format
    metadata["generated_code"] = cv_code
    metadata["generated_text"] = cv_code
    return metadata


def get_applicant_cv(email: str) -> dict | None:
    try:
        response = supabase.table(CV_METADATA_TABLE).select("*").eq("applicant_email", email).execute()
        if response.data:
            record = response.data[0]
            raw_profile_data = record.get("profile_data")
            normalized_profile_data = normalize_profile_data(raw_profile_data)
            record["profile_data"] = normalized_profile_data
            return record
    except Exception as ex:
        return None
    return None

# ---- estilos globales para una apariencia minimalista ----
UI_SPACING = {
    "xs": 4,
    "sm": 8,
    "md": 12,
    "lg": 16,
    "xl": 20,
}

UI_RADIUS = {
    "sm": 8,
    "md": 12,
    "lg": 16,
    "xl": 22,
}

UI_TYPE_SCALE = {
    "title_large": {"mobile": 28, "desktop": 34},
    "title": {"mobile": 22, "desktop": 28},
    "subtitle": {"mobile": 16, "desktop": 18},
    "body": {"mobile": 13, "desktop": 14},
    "label": {"mobile": 11, "desktop": 12},
}

BASE_BUTTON_STYLE = ft.ButtonStyle(
    shape=ft.RoundedRectangleBorder(radius=UI_RADIUS["md"]),
    text_style=ft.TextStyle(size=14, weight=ft.FontWeight.W_600),
)

# helper para crear botones consistentes

def mk_button(
    label: str,
    on_click=None,
    style: ft.ButtonStyle | None = None,
    variant: str = "secondary",
    palette: dict[str, str] | None = None,
    **kwargs,
) -> ft.Button:
    if style is None and palette is not None:
        style = get_button_style(variant, palette)
    return ft.Button(label, on_click=on_click, style=style or BASE_BUTTON_STYLE, **kwargs)

# helper para cajas de texto con diseño limpio

def mk_textfield(label: str, palette: dict[str, str] | None = None, **kwargs) -> ft.TextField:
    field_kwargs = {
        "label": label,
        "border_radius": UI_RADIUS["sm"],
    }
    if palette is not None:
        field_kwargs.update(
            {
                "filled": True,
                "bgcolor": palette["surface_soft"],
                "border_color": palette["border"],
                "focused_border_color": palette["accent"],
                "color": palette["text"],
                "label_style": ft.TextStyle(color=palette["muted"]),
                "cursor_color": palette["accent"],
            }
        )
    field_kwargs.update(kwargs)
    return ft.TextField(**field_kwargs)


def is_dark_mode(page: ft.Page) -> bool:
    return page.theme_mode == ft.ThemeMode.DARK


def get_palette(page: ft.Page) -> dict[str, str]:
    dark = is_dark_mode(page)
    if dark:
        return {
            "bg": "#050505",
            "surface": "#0D0D0D",
            "surface_soft": "#141414",
            "border": "#222222",
            "text": "#F5F5F5",
            "muted": "#9A9AA0",
            "accent": "#F5F5F5",
            "accent_text": "#0B0B0B",
            "danger": "#EF4444",
            "success": "#22C55E",
            "warning": "#F59E0B",
            "info": "#38BDF8",
        }
    return {
        "bg": "#F3F6FB",
        "surface": "#FFFFFF",
        "surface_soft": "#F8FAFC",
        "border": "#DDE3EE",
        "text": "#0F172A",
        "muted": "#64748B",
        "accent": "#0F172A",
        "accent_text": "#FFFFFF",
        "danger": "#DC2626",
        "success": "#16A34A",
        "warning": "#D97706",
        "info": "#0284C7",
    }


def get_button_style(variant: str, palette: dict[str, str]) -> ft.ButtonStyle:
    variants = {
        "primary": {
            "bgcolor": palette["accent"],
            "color": palette["accent_text"],
            "side": ft.BorderSide(1, palette["accent"]),
        },
        "secondary": {
            "bgcolor": palette["surface_soft"],
            "color": palette["text"],
            "side": ft.BorderSide(1, palette["border"]),
        },
        "ghost": {
            "bgcolor": palette["surface"],
            "color": palette["text"],
            "side": ft.BorderSide(1, palette["surface"]),
        },
        "danger": {
            "bgcolor": palette["surface_soft"],
            "color": palette["danger"],
            "side": ft.BorderSide(1, palette["border"]),
        },
    }
    selected = variants.get(variant, variants["secondary"])
    return ft.ButtonStyle(
        bgcolor=selected["bgcolor"],
        color=selected["color"],
        side=selected["side"],
        shape=ft.RoundedRectangleBorder(radius=UI_RADIUS["md"]),
        text_style=ft.TextStyle(size=14, weight=ft.FontWeight.W_700),
    )


def ui_font_size(page: ft.Page, token: str) -> int:
    scale = UI_TYPE_SCALE.get(token, UI_TYPE_SCALE["body"])
    return scale["mobile"] if is_mobile(page) else scale["desktop"]


def get_page_padding(page: ft.Page) -> int:
    if is_mobile(page):
        return UI_SPACING["sm"]
    if is_tablet(page):
        return UI_SPACING["md"]
    return UI_SPACING["xl"]


def show_feedback(page: ft.Page, message: str, tone: str = "info", action_label: str | None = None, on_action=None):
    palette = get_palette(page)
    color_by_tone = {
        "success": palette["success"],
        "error": palette["danger"],
        "warning": palette["warning"],
        "info": palette["info"],
    }
    bgcolor = color_by_tone.get(tone, palette["info"])
    snack = ft.SnackBar(
        content=ft.Text(message, color="#FFFFFF", size=12),
        bgcolor=bgcolor,
    )
    if action_label and on_action:
        snack.action = action_label
        snack.on_action = on_action
    page.snack_bar = snack
    page.snack_bar.open = True
    page.update()


def mk_surface_card(
    page: ft.Page,
    palette: dict[str, str],
    content,
    *,
    expand: bool = False,
    padding: int | None = None,
    variant: str = "default",
):
    card_padding = padding if padding is not None else (UI_SPACING["md"] if is_mobile(page) else UI_SPACING["lg"])
    card_radius = UI_RADIUS["lg"] if is_mobile(page) else UI_RADIUS["xl"]
    bgcolor = palette["surface"]
    border_color = palette["border"]
    if variant == "hero":
        bgcolor = palette["surface"]
        border_color = palette["accent"]
    elif variant == "secondary":
        bgcolor = palette["surface_soft"]
    return ft.Container(
        bgcolor=bgcolor,
        border=ft.Border.all(1, border_color),
        border_radius=card_radius,
        padding=card_padding,
        content=content,
        expand=expand,
    )


def get_choice_button_style(selected: bool, palette: dict[str, str]) -> ft.ButtonStyle:
    if selected:
        return ft.ButtonStyle(
            bgcolor=palette["accent"],
            color=palette["accent_text"],
            side=ft.BorderSide(1, palette["accent"]),
            shape=ft.RoundedRectangleBorder(radius=UI_RADIUS["sm"]),
            text_style=ft.TextStyle(weight=ft.FontWeight.W_700),
        )
    return ft.ButtonStyle(
        bgcolor=palette["surface_soft"],
        color=palette["text"],
        side=ft.BorderSide(1, palette["border"]),
        shape=ft.RoundedRectangleBorder(radius=UI_RADIUS["sm"]),
    )


def mk_section_header(page: ft.Page, palette: dict[str, str], title: str, subtitle: str = ""):
    controls = [
        ft.Text(title, weight=ft.FontWeight.W_900, size=ui_font_size(page, "title"), color=palette["text"]),
    ]
    if subtitle:
        controls.append(ft.Text(subtitle, size=ui_font_size(page, "body"), color=palette["muted"]))
    return ft.Column(controls, spacing=UI_SPACING["xs"])


def mk_status_chip(page: ft.Page, palette: dict[str, str], label: str, count: int, color: str):
    return ft.Container(
        bgcolor=palette["surface_soft"],
        border=ft.Border.all(1, palette["border"]),
        border_radius=UI_RADIUS["md"],
        padding=UI_SPACING["sm"],
        content=ft.Row(
            [
                ft.Container(width=10, height=10, border_radius=5, bgcolor=color),
                ft.Text(label, size=ui_font_size(page, "label"), color=palette["text"], expand=True),
                ft.Text(str(count), size=ui_font_size(page, "label"), weight=ft.FontWeight.W_700, color=palette["text"]),
            ],
            spacing=UI_SPACING["sm"],
            alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
        ),
    )


def mk_chat_bubble(
    page: ft.Page,
    palette: dict[str, str],
    content,
    *,
    role: str,
    width: int | None = None,
):
    is_user = role == "user"
    return ft.Container(
        content=content,
        bgcolor=palette["surface_soft"] if is_user else palette["accent"],
        border_radius=UI_RADIUS["md"],
        padding=UI_SPACING["sm"],
        margin=ft.Margin(left=0 if is_user else 60, top=4, right=60 if is_user else 0, bottom=4),
        width=width,
    )


def mk_centered_shell(page: ft.Page, palette: dict[str, str], content, *, width: int | None = None):
    return ft.Container(
        expand=True,
        alignment=ft.Alignment.CENTER,
        content=mk_surface_card(
            page,
            palette,
            content,
            variant="hero",
            padding=18 if is_mobile(page) else 24 if is_tablet(page) else 30,
        ) if width is None else ft.Container(
            width=width,
            content=mk_surface_card(
                page,
                palette,
                content,
                variant="hero",
                padding=18 if is_mobile(page) else 24 if is_tablet(page) else 30,
            ),
        ),
    )


def mk_loading_overlay(
    page: ft.Page,
    palette: dict[str, str],
    title: str,
    subtitle: str = "",
    *,
    visible: bool = False,
):
    return ft.Container(
        visible=visible,
        alignment=ft.Alignment.CENTER,
        bgcolor="rgba(0, 0, 0, 0.50)",
        left=0,
        top=0,
        right=0,
        bottom=0,
        expand=True,
        content=ft.Container(
            width=280 if is_mobile(page) else 320,
            padding=UI_SPACING["lg"],
            border_radius=UI_RADIUS["lg"],
            bgcolor=palette["surface"],
            content=ft.Column(
                [
                    ft.ProgressRing(width=46, height=46, stroke_width=4, color=palette["accent"]),
                    ft.Text(title, size=ui_font_size(page, "subtitle"), weight=ft.FontWeight.W_700, color=palette["text"], text_align=ft.TextAlign.CENTER),
                    ft.Text(subtitle, size=ui_font_size(page, "label"), color=palette["muted"], text_align=ft.TextAlign.CENTER),
                ],
                spacing=UI_SPACING["sm"],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                alignment=ft.MainAxisAlignment.CENTER,
                tight=True,
            ),
        ),
    )


def mk_empty_state(page: ft.Page, palette: dict[str, str], title: str, subtitle: str = ""):
    return ft.Column(
        [
            ft.Icon(ft.Icons.INBOX_OUTLINED, size=32 if is_mobile(page) else 40, color=palette["muted"]),
            ft.Text(title, size=ui_font_size(page, "subtitle"), weight=ft.FontWeight.W_700, color=palette["text"], text_align=ft.TextAlign.CENTER),
            ft.Text(subtitle, size=ui_font_size(page, "body"), color=palette["muted"], text_align=ft.TextAlign.CENTER),
        ],
        spacing=UI_SPACING["xs"],
        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
        alignment=ft.MainAxisAlignment.CENTER,
    )


def mk_status_text(page: ft.Page, palette: dict[str, str], text: str = "", tone: str = "muted"):
    color_map = {
        "muted": palette["muted"],
        "error": palette["danger"],
        "success": palette["success"],
        "warning": palette["warning"],
        "info": palette["info"],
    }
    return ft.Text(text, color=color_map.get(tone, palette["muted"]), size=ui_font_size(page, "label"))


def get_nav_icon_button_style(is_active: bool, palette: dict[str, str]) -> ft.ButtonStyle:
    return ft.ButtonStyle(
        bgcolor=palette["accent"] if is_active else palette["surface_soft"],
        shape=ft.RoundedRectangleBorder(radius=UI_RADIUS["lg"]),
    )


def get_icon_button_surface_style(palette: dict[str, str], active: bool = False) -> ft.ButtonStyle:
    return ft.ButtonStyle(
        bgcolor=palette["accent"] if active else palette["surface"],
        side=ft.BorderSide(1, palette["accent"] if active else palette["border"]),
        shape=ft.RoundedRectangleBorder(radius=UI_RADIUS["lg"]),
    )


def make_theme_toggle(page: ft.Page, on_click) -> ft.IconButton:
    palette = get_palette(page)
    return ft.IconButton(
        icon=ft.Icons.LIGHT_MODE if is_dark_mode(page) else ft.Icons.DARK_MODE,
        tooltip="Cambiar tema",
        on_click=on_click,
        icon_color=palette["text"],
        style=get_icon_button_surface_style(palette),
    )


def get_viewport_width(page: ft.Page) -> int:
    width = page.width or 0
    if not width:
        try:
            width = page.window.width or 0
        except Exception:
            width = 0
    return int(width or 1200)


def is_mobile(page: ft.Page) -> bool:
    return get_viewport_width(page) < 768


def is_tablet(page: ft.Page) -> bool:
    width = get_viewport_width(page)
    return 768 <= width < 1100


def get_device_bucket(page: ft.Page) -> str:
    if is_mobile(page):
        return "mobile"
    if is_tablet(page):
        return "tablet"
    return "desktop"


def make_responsive_resize_handler(page: ft.Page, rebuild_view):
    state = {"device": get_device_bucket(page)}

    def _on_resized(e=None):
        new_device = get_device_bucket(page)
        if new_device != state["device"]:
            state["device"] = new_device
            asyncio.create_task(rebuild_view())
            return
        page.update()

    return _on_resized


def normalize_text(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", text).strip()


def trim_text(text: str, max_len: int = 340) -> str:
    text = normalize_text(text)
    if len(text) <= max_len:
        return text
    return text[: max_len - 1].rstrip() + "…"


APPLICATION_STATUS_OPTIONS = ["Aplicado", "En proceso", "En revisión", "Rechazado"]
APPLICATION_STATUS_COLORS = {
    "Aplicado": "#2563EB",
    "En proceso": "#7C3AED",
    "En revisión": "#F59E0B",
    "Rechazado": "#EF4444",
}


def get_application_title(job: dict) -> str:
    return normalize_text(job.get("vaccancy", "") or job.get("vacancy", "") or "Sin cargo")


def get_application_description(job: dict) -> str:
    return normalize_text(job.get("Description", "") or job.get("description", "") or "")


def build_application_description_payload(job_data: dict, enriched_description: str) -> str:
    title = normalize_text(job_data.get("title", ""))
    company = normalize_text(job_data.get("company", ""))
    location = normalize_text(job_data.get("location", ""))
    website = normalize_text(job_data.get("website", ""))

    lines = [
        f"Cargo: {title or 'Sin cargo'}",
        f"Empresa/Sitio: {company or website or 'No disponible'}",
        f"Ubicación: {location or 'No disponible'}",
        "",
        enriched_description or "Sin descripción enriquecida.",
    ]
    return "\n".join(lines)


def build_interview_system_prompt(application_row: dict) -> str:
    job_title = get_application_title(application_row)
    job_desc = get_application_description(application_row)
    link = normalize_text(application_row.get("application_link", ""))

    return (
        "Actua como coach experto en entrevistas laborales. "
        "Tu objetivo es preparar al usuario para una entrevista real de esta vacante.\n\n"
        f"Vacante: {job_title}\n"
        f"Enlace: {link or 'No disponible'}\n"
        "Contexto de vacante (Description):\n"
        f"{job_desc or 'No disponible'}\n\n"
        "Reglas:\n"
        "1) Responde en español claro y profesional.\n"
        "2) Da consejos prácticos, preguntas probables y ejemplos de respuesta.\n"
        "3) Si falta contexto, pide información puntual al usuario.\n"
        "4) Mantén respuestas accionables y enfocadas al rol específico."
    )


def generate_interview_reply_with_gemini(application_row: dict, history: list[dict[str, str]], user_message: str) -> str:
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY no configurada.")

    conversation_parts: list[str] = []
    for msg in history[-12:]:
        role = normalize_text(msg.get("role", ""))
        content = normalize_text(msg.get("content", ""))
        if role and content:
            speaker = "Usuario" if role == "user" else "Asistente"
            conversation_parts.append(f"{speaker}: {content}")

    conversation_parts.append(f"Usuario: {normalize_text(user_message)}")
    conversation_text = "\n".join(conversation_parts)

    prompt = (
        f"{build_interview_system_prompt(application_row)}\n\n"
        "Historial de conversación:\n"
        f"{conversation_text}\n\n"
        "Responde como Asistente:"
    )

    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.6,
            "topP": 0.9,
            "maxOutputTokens": 1024,
        },
    }

    models_to_try = [GEMINI_MODEL] + [m for m in GEMINI_FALLBACK_MODELS if m != GEMINI_MODEL]
    last_status_code = 0
    last_error_detail = ""

    for model_name in models_to_try:
        url = f"https://generativelanguage.googleapis.com/{GEMINI_API_VERSION}/models/{model_name}:generateContent"
        try:
            response = httpx.post(
                url,
                json=payload,
                headers={"x-goog-api-key": GEMINI_API_KEY},
                timeout=45.0,
            )
            response.raise_for_status()
            gemini_text = extract_gemini_text(response.json())
            if gemini_text:
                return gemini_text
            raise ValueError("Gemini no devolvio contenido para entrevista.")
        except httpx.HTTPStatusError as ex:
            last_status_code = ex.response.status_code
            last_error_detail = extract_gemini_error_detail(ex.response)
            if ex.response.status_code == 404:
                continue
            if ex.response.status_code == 429:
                raise ValueError("Gemini devolvio 429 por limite de cuota o solicitudes.") from None
            detail = f": {last_error_detail}" if last_error_detail else ""
            raise ValueError(f"Gemini devolvio error HTTP {ex.response.status_code}{detail}") from None
        except httpx.HTTPError as ex:
            raise ValueError(f"No fue posible conectar con Gemini: {ex}") from None

    if last_status_code == 404:
        raise ValueError("Gemini devolvio 404 para los modelos configurados.")
    raise ValueError("No fue posible generar respuesta de entrevista con Gemini.")


def collect_highlight_items(job: dict) -> list[str]:
    items: list[str] = []
    for block in job.get("job_highlights", []) or []:
        title = normalize_text(block.get("title", ""))
        values = block.get("items", []) or []
        values = [normalize_text(v) for v in values if normalize_text(v)]
        if title and values:
            items.append(f"{title}: {', '.join(values)}")
        elif values:
            items.append(", ".join(values))
    return items


def fetch_page_description(url: str, timeout: int = 6) -> str:
    if not url or url == "#":
        return ""

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        )
    }

    try:
        response = requests.get(url, timeout=timeout, headers=headers)
        response.raise_for_status()
    except Exception:
        return ""

    soup = BeautifulSoup(response.text, "html.parser")
    candidates: list[str] = []

    for attrs in (
        {"name": "description"},
        {"property": "og:description"},
        {"name": "twitter:description"},
    ):
        meta = soup.find("meta", attrs=attrs)
        content = normalize_text(meta.get("content", "")) if meta else ""
        if content and len(content) > 40:
            candidates.append(content)

    h1 = soup.find("h1")
    if h1:
        heading = normalize_text(h1.get_text(" ", strip=True))
        if heading and len(heading) > 10:
            candidates.append(heading)

    paragraphs = []
    for p in soup.find_all("p", limit=20):
        t = normalize_text(p.get_text(" ", strip=True))
        if len(t) >= 70:
            paragraphs.append(t)
        if len(paragraphs) >= 2:
            break
    candidates.extend(paragraphs)

    # Quitar repetidos preservando orden.
    unique: list[str] = []
    seen = set()
    for c in candidates:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            unique.append(c)

    return trim_text(" ".join(unique), max_len=420)


def build_enriched_description(job: dict, apply_link: str) -> str:
    title = normalize_text(job.get("title", ""))
    company = normalize_text(job.get("company_name", ""))
    location = normalize_text(job.get("location", ""))
    posted = normalize_text(job.get("detected_extensions", {}).get("posted_at", ""))

    serp_desc = normalize_text(job.get("description", ""))
    highlights = collect_highlight_items(job)

    # Intentamos enriquecer con información de la página de detalle/aplicación.
    page_desc = fetch_page_description(job.get("share_link") or apply_link)

    description_blocks: list[str] = []
    if serp_desc:
        description_blocks.append(serp_desc)
    if highlights:
        description_blocks.append("\n".join([f"- {item}" for item in highlights]))
    if page_desc:
        description_blocks.append(page_desc)

    sections = [
        f"Cargo: {title or 'Sin título'}",
        f"Empresa: {company or 'N/A'}",
        f"Ubicación: {location or 'N/A'}",
        f"Fecha: {posted or 'N/A'}",
        "",
        "Descripción:",
        "\n\n".join(description_blocks) if description_blocks else "Sin descripción disponible.",
    ]

    return "\n".join(sections).strip()


async def logout_user(page: ft.Page):
    global current_user
    current_user = None
    user_details.clear()
    remove_saved_credentials()
    await show_login_ui(page)


async def show_add_manual_job_dialog(page: ft.Page):
    """Función para agregar manualmente un empleo por URL."""
    palette = get_palette(page)
    add_job_state = {"processing": False}

    url_input = mk_textfield("URL del empleo", expand=True)
    message = ft.Text("", color=ft.Colors.RED)
    add_button = mk_button("Agregar Empleo")
    add_button.disabled = True

    def validate_url(evt=None):
        url = url_input.value.strip()
        is_valid = url.startswith("http://") or url.startswith("https://")
        add_button.disabled = not is_valid or add_job_state["processing"]
        page.update()

    url_input.on_change = validate_url

    def close_dialog(dlg):
        dlg.open = False
        page.update()

    async def process_add_job(evt):
        if add_job_state["processing"]:
            return

        add_job_state["processing"] = True
        add_button.disabled = True
        message.value = "Procesando..."
        message.color = ft.Colors.BLUE
        page.update()

        url = url_input.value.strip()

        try:
            if not url.startswith("http://") and not url.startswith("https://"):
                message.value = "Por favor ingresa una URL válida"
                add_job_state["processing"] = False
                validate_url()
                page.update()
                return

            resp = requests.get(url, timeout=5)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "html.parser")

            titulo = soup.find("title").text.strip() if soup.find("title") else "Empleo sin título"

            meta = soup.find("meta", attrs={"name": "description"}) or \
                   soup.find("meta", attrs={"property": "og:description"})
            descripcion = meta["content"].strip() if meta and meta.get("content") else "Sin descripción"

            parsed = urlparse(url)
            sitio = parsed.netloc or ""
            sitio = (parsed.netloc[4:] if parsed.netloc.startswith("www.") else parsed.netloc) or ""

            supabase.table("Applications").insert({
                "applicant_email": current_user,
                "website": sitio,
                "vaccancy": titulo,
                "status": "Aplicado",
                "application_link": url,
                "Description": descripcion,
            }).execute()

            page.snack_bar = ft.SnackBar(
                content=ft.Text("Empleo agregado exitosamente"),
                bgcolor=ft.Colors.GREEN
            )
            page.snack_bar.open = True

            dlg.open = False
            page.update()
            await show_menu_ui(page)

        except Exception as ex:
            add_job_state["processing"] = False
            message.color = ft.Colors.RED

            if "Timeout" in str(type(ex)):
                message.value = "La URL tardó demasiado en responder."
            else:
                message.value = "Error al guardar. Verfica la URL o tu conexión."

            validate_url()
            page.update()

    add_button.on_click = process_add_job

    dlg = ft.AlertDialog(
        title=ft.Text("Agregar Empleo Manualmente"),
        content=ft.Column([
            url_input,
            message,
        ], spacing=12, tight=True),
        actions=[
            add_button,
            ft.TextButton("Cancelar", on_click=lambda _: close_dialog(dlg)),
        ],
    )

    page.overlay.append(dlg)
    dlg.open = True
    page.update()


def open_ai_cv_builder_dialog(
    page: ft.Page,
    palette: dict[str, str],
    initial_values: dict[str, str] | None,
    on_complete,
    generate_and_store: bool = False,
):
    values = normalize_profile_data(initial_values)
    dlg: ft.AlertDialog | None = None

    design_steps = [
        {
            "key": "cv_color_palette",
            "prompt": "Elige la paleta de colores para tu CV:",
            "options": AI_CV_COLOR_PALETTES,
            "normalize": normalize_ai_cv_color_palette,
            "label": get_ai_cv_color_palette_label,
        },
        {
            "key": "cv_font_size",
            "prompt": "Ahora selecciona el tamaño de letra:",
            "options": AI_CV_FONT_SIZE_OPTIONS,
            "normalize": normalize_ai_cv_font_size,
            "label": get_ai_cv_font_size_label,
        },
        {
            "key": "cv_columns",
            "prompt": "Selecciona la distribución de columnas:",
            "options": AI_CV_COLUMN_OPTIONS,
            "normalize": normalize_ai_cv_columns,
            "label": get_ai_cv_columns_label,
        },
        {
            "key": "cv_include_photo",
            "prompt": "Finalmente, ¿quieres incluir foto en el CV?",
            "options": AI_CV_PHOTO_OPTIONS,
            "normalize": normalize_ai_cv_photo_option,
            "label": get_ai_cv_photo_option_label,
        },
    ]

    output_format_options = [
        {
            "id": "docx",
            "name": "DOCX",
            "description": "Editable",
        },
        {
            "id": "pdf",
            "name": "PDF",
            "description": "No editable",
        },
    ]

    # Estado del chat
    chat_state = {
        "current_section": 0,
        "processing": False,
        "awaiting_choice": None,
        "design_step_index": 0,
        "output_format": normalize_ai_cv_output_format(values.get("output_format", "docx")),
        "profile_data": {
            "target_roles": values.get("target_roles", ""),
            "summary": values.get("summary", ""),
            "skills": values.get("skills", ""),
            "experience": values.get("experience", ""),
            "education": values.get("education", ""),
            "languages": values.get("languages", ""),
            "certifications": values.get("certifications", ""),
            "achievements": values.get("achievements", ""),
            "cv_color_palette": normalize_ai_cv_color_palette(values.get("cv_color_palette", "")),
            "cv_font_size": normalize_ai_cv_font_size(values.get("cv_font_size", "")),
            "cv_columns": normalize_ai_cv_columns(values.get("cv_columns", "")),
            "cv_include_photo": normalize_ai_cv_photo_option(values.get("cv_include_photo", "")),
            "cv_photo_base64": values.get("cv_photo_base64", ""),
        },
    }
    
    # Configuración de secciones
    sections = [
        {
            "key": "target_roles",
            "label": "Vacantes Objetivo",
            "prompt": "¿A qué vacantes o posiciones te gustaría aplicar? (ej: Ingeniero de Software, Analista QA, Data Scientist)",
            "example": "Ingeniero Senior de Software, Full Stack Developer",
            "min_length": 5,
        },
        {
            "key": "summary",
            "label": "Resumen Profesional",
            "prompt": "Por favor, comparte un resumen breve de tu perfil profesional (experiencia clave, fortalezas).",
            "example": "Ingeniero de software con 5 años de experiencia en desarrollo web...",
            "min_length": 20,
            "required": True,
        },
        {
            "key": "skills",
            "label": "Habilidades",
            "prompt": "¿Cuáles son tus habilidades técnicas principales? (puedes separarlas por comas)",
            "example": "Python, JavaScript, React, Django, SQL, Git",
            "min_length": 5,
        },
        {
            "key": "experience",
            "label": "Experiencia",
            "prompt": "Cuéntame sobre tu experiencia laboral. Incluye cargos, empresas y logros principales.",
            "example": "2018-2020: Desarrollador Junior en TechCorp | 2020-Presente: Senior Dev en StartupXYZ",
            "min_length": 20,
            "required": True,
        },
        {
            "key": "education",
            "label": "Educación",
            "prompt": "¿Cuál es tu formación académica? (grados, institutos, certificaciones relevantes)",
            "example": "Licenciatura en Ingeniería de Sistemas - Universidad Nacional",
            "min_length": 15,
            "required": True,
        },
        {
            "key": "languages",
            "label": "Idiomas",
            "prompt": "¿Qué idiomas hablas? (especifica nivel: básico, intermedio, avanzado, nativo)",
            "example": "Español (nativo), Inglés (avanzado), Francés (intermedio)",
            "min_length": 3,
        },
        {
            "key": "certifications",
            "label": "Certificaciones",
            "prompt": "¿Tienes certificaciones relevantes? (pueden ser profesionales, tecnológicas, etc.)",
            "example": "AWS Solutions Architect, Google Cloud Associate, Scrum Master",
            "min_length": 3,
        },
        {
            "key": "achievements",
            "label": "Logros",
            "prompt": "¿Cuáles han sido tus logros más importantes en tu carrera profesional?",
            "example": "Lideré proyecto que redujo costos 40%, Mentorizo a 3 desarrolladores junior",
            "min_length": 10,
        },
    ]
    
    # Componentes de UI
    messages_list = ft.ListView(spacing=8, expand=True, auto_scroll=True)
    result_actions = ft.Row([], wrap=True, spacing=8)
    choice_quick_actions = ft.Row([], wrap=True, spacing=8, visible=False)
    status_text = mk_status_text(page, palette)
    input_field = mk_textfield("Tu respuesta", palette=palette, multiline=False, min_lines=1)

    send_button = ft.IconButton(
        icon=ft.Icons.SEND,
        icon_color=palette["accent"],
        tooltip="Enviar respuesta",
    )

    photo_picker = ft.FilePicker()

    def ensure_photo_picker():
        if photo_picker not in page.overlay:
            page.overlay.append(photo_picker)

    def run_async(coro):
        task = asyncio.create_task(coro)

        def _consume_task_error(done_task: asyncio.Task):
            try:
                task_error = done_task.exception()
                if task_error:
                    error_name = type(task_error).__name__
                    error_msg = str(task_error)
                    # Solo imprimir si NO es el error conocido de focus en AlertDialog
                    if not ("Control must be added to the page first" in error_msg and "focus" in error_msg.lower()):
                        if not ("Control must be added" in error_msg):
                            pass
            except Exception:
                return

        task.add_done_callback(_consume_task_error)

    def open_dialog(dialog: ft.AlertDialog):
        if dialog not in page.overlay:
            page.overlay.append(dialog)
        
        dialog.open = True
        
        page.update()

    def add_message(role: str, content: str):
        try:
            if role == "bot":
                msg_bubble = mk_chat_bubble(
                    page,
                    palette,
                    content=ft.Text(
                        content, 
                        color=palette["bg"], 
                        size=12, 
                        selectable=True,
                        text_align=ft.TextAlign.LEFT,
                    ),
                    role="bot",
                )
            else:  # user
                msg_bubble = mk_chat_bubble(
                    page,
                    palette,
                    content=ft.Text(
                        content,
                        color=palette["text"],
                        size=12,
                        selectable=True,
                        text_align=ft.TextAlign.RIGHT,
                    ),
                    role="user",
                )
            
            messages_list.controls.append(msg_bubble)
        except Exception:
            return

    def close_dialog(e=None):
        nonlocal dlg
        if dlg is None:
            return
        # Patrón estable en Flet: cerrar el diálogo sin removerlo manualmente del overlay.
        dlg.open = False
        page.update()

    def set_input_enabled(enabled: bool):
        input_field.disabled = not enabled
        send_button.disabled = not enabled

    def set_choice_mode(active: bool):
        set_input_enabled(not active)
        choice_quick_actions.visible = active

    def ask_current_question():
        current_idx = chat_state["current_section"]
        if current_idx >= len(sections):
            return
        set_choice_mode(False)
        section = sections[current_idx]
        add_message("bot", f"{section['label']}\n\n{section['prompt']}")

    async def handle_choice_selection(selected_id: str):
        if chat_state["processing"]:
            return

        awaiting = chat_state.get("awaiting_choice")
        if not awaiting:
            return

        if awaiting == "output_format":
            chat_state["output_format"] = normalize_ai_cv_output_format(selected_id)
            chat_state["profile_data"]["output_format"] = chat_state["output_format"]
            add_message("user", chat_state["output_format"].upper())
            choice_quick_actions.controls.clear()
            set_choice_mode(False)
            chat_state["awaiting_choice"] = None
            add_message("bot", f"Perfecto. Usaremos formato {chat_state['output_format'].upper()}.")
            await finalize_flow()
            page.update()
            return

        if awaiting == "photo_upload":
            if selected_id == "photo_pick":
                ensure_photo_picker()
                photo_picker.pick_files(
                    allow_multiple=False,
                    allowed_extensions=["png", "jpg", "jpeg"],
                    dialog_title="Selecciona tu foto de perfil",
                )
                return

            if selected_id == "photo_continue":
                has_photo = bool(chat_state["profile_data"].get("cv_photo_base64", ""))
                if not has_photo:
                    add_message("bot", "Primero selecciona una foto para poder continuar.")
                    page.update()
                    return
                add_message("user", "Continuar")
                choice_quick_actions.controls.clear()
                set_choice_mode(False)
                chat_state["awaiting_choice"] = None
                ask_output_format_question()
                page.update()
                return

            if selected_id == "photo_skip":
                add_message("user", "Continuar sin foto")
                chat_state["profile_data"]["cv_include_photo"] = "sin_foto"
                chat_state["profile_data"]["cv_photo_base64"] = ""
                choice_quick_actions.controls.clear()
                set_choice_mode(False)
                chat_state["awaiting_choice"] = None
                ask_output_format_question()
                page.update()
                return

        current_step = design_steps[chat_state["design_step_index"]]
        normalized_value = current_step["normalize"](selected_id)
        selected_label = current_step["label"](normalized_value)

        chat_state["profile_data"][current_step["key"]] = normalized_value
        add_message("user", selected_label)

        if current_step["key"] == "cv_include_photo" and normalized_value == "con_foto":
            chat_state["design_step_index"] += 1
            ask_photo_upload_question()
            page.update()
            return

        if current_step["key"] == "cv_include_photo" and normalized_value == "sin_foto":
            chat_state["profile_data"]["cv_photo_base64"] = ""

        chat_state["design_step_index"] += 1
        ask_next_design_choice()
        page.update()

    def build_choice_button(option: dict[str, str], selected_id: str | None = None) -> ft.Button:
        is_selected = selected_id == option["id"]
        label = f"{option['name']} ✓" if is_selected else option["name"]
        return mk_button(
            label,
            on_click=lambda e, sid=option["id"]: run_async(handle_choice_selection(sid)),
            style=get_choice_button_style(is_selected, palette),
        )

    def render_choice_buttons(options: list[dict[str, str]], selected_id: str | None = None):
        choice_quick_actions.controls = [build_choice_button(option, selected_id) for option in options]
        set_choice_mode(True)

    def ask_output_format_question():
        chat_state["awaiting_choice"] = "output_format"
        render_choice_buttons(output_format_options, selected_id=chat_state.get("output_format"))
        add_message(
            "bot",
            "Elige el formato final del archivo:",
        )

    def ask_photo_upload_question():
        chat_state["awaiting_choice"] = "photo_upload"
        has_photo = bool(chat_state["profile_data"].get("cv_photo_base64", ""))
        photo_options = [
            {"id": "photo_pick", "name": "Seleccionar foto", "description": "PNG/JPG/JPEG"},
            {"id": "photo_continue", "name": "Continuar", "description": "Usar foto seleccionada"},
            {"id": "photo_skip", "name": "Continuar sin foto", "description": "Desactivar foto"},
        ]
        render_choice_buttons(photo_options)
        if has_photo:
            add_message("bot", "Foto cargada correctamente. Puedes continuar o reemplazarla.")
        else:
            add_message("bot", "Selecciona una foto real de perfil para incrustarla en el CV.")

    def ask_next_design_choice():
        if chat_state["design_step_index"] >= len(design_steps):
            ask_output_format_question()
            return

        step = design_steps[chat_state["design_step_index"]]
        current_value = step["normalize"](chat_state["profile_data"].get(step["key"], ""))
        chat_state["profile_data"][step["key"]] = current_value
        chat_state["awaiting_choice"] = step["key"]

        lines = [step["prompt"]]
        for option in step["options"]:
            lines.append(f"- {option['name']}: {option['description']}")
        add_message("bot", "\n".join(lines))
        render_choice_buttons(step["options"], selected_id=current_value)

    async def generate_cv_html_from_profile(
        profile_data: dict,
        full_name: str,
        email: str,
        output_format: str,
    ) -> str:
        """Envía un prompt completo a Gemini solicitando el HTML del CV y devuelve la cadena HTML."""
        style_profiles = {
            "azul_profesional": {
                "style_name": "Corporativo Azul",
                "header_primary": "#1e3a8a",
                "accent": "#3b82f6",
                "text": "#0f172a",
                "muted": "#475569",
                "sidebar_bg": "#f8fafc",
                "page_bg": "#eff6ff",
            },
            "verde_moderno": {
                "style_name": "Moderno Verde",
                "header_primary": "#166534",
                "accent": "#16a34a",
                "text": "#0f172a",
                "muted": "#475569",
                "sidebar_bg": "#f0fdf4",
                "page_bg": "#f7fee7",
            },
            "gris_ejecutivo": {
                "style_name": "Ejecutivo Gris",
                "header_primary": "#374151",
                "accent": "#4b5563",
                "text": "#111827",
                "muted": "#4b5563",
                "sidebar_bg": "#f3f4f6",
                "page_bg": "#f9fafb",
            },
        }
        font_pt_map = {"compacta": "10pt", "estandar": "11pt", "amplia": "13pt"}
        columns_desc_map = {
            "una_columna": "una columna principal (sin barra lateral)",
            "dos_columnas": "dos columnas: izquierda 2/3 y derecha 1/3",
        }

        selected_palette_id = normalize_ai_cv_color_palette(profile_data.get("cv_color_palette", ""))
        selected_font_id    = normalize_ai_cv_font_size(profile_data.get("cv_font_size", ""))
        selected_columns_id = normalize_ai_cv_columns(profile_data.get("cv_columns", ""))
        include_photo       = normalize_ai_cv_photo_option(profile_data.get("cv_include_photo", "")) == "con_foto"

        style_cfg   = style_profiles.get(selected_palette_id, style_profiles["azul_profesional"])
        font_size   = font_pt_map.get(selected_font_id, "11pt")
        layout_desc = columns_desc_map.get(selected_columns_id, "una sola columna")

        photo_instruction = (
            'Incluye un elemento <img id="cv-photo" alt="Foto de perfil" '
            'style="width:110px;height:130px;object-fit:cover;border-radius:4px;'
            'float:right;margin:0 0 12px 16px;" />'
        ) if include_photo else "No incluyas elemento de foto."

        column_instruction = (
            "Implementa el layout de dos columnas solicitado (izquierda 2/3 para resumen, experiencia y educación; "
            "derecha 1/3 para habilidades, idiomas, certificaciones y logros)."
            if selected_columns_id == "dos_columnas"
            else "Implementa una sola columna legible y amplia, manteniendo secciones bien separadas."
        )

        pdf_compat_instruction = (
            "Como el formato final es PDF, evita pseudo-elementos CSS (::before/::after) y evita dependencias avanzadas "
            "de flex/grid que puedan romper convertidores HTML->PDF; usa HTML/CSS compatible."
            if normalize_ai_cv_output_format(output_format) == "pdf"
            else "Puedes usar recursos modernos de CSS manteniendo compatibilidad básica de renderizado."
        )

        html_prompt = (
            "Actúa como un desarrollador Frontend experto en diseño de documentos profesionales. "
            "Genera un archivo HTML5 y CSS3 para un currículum vitae (CV) que siga estas especificaciones estrictas.\n\n"
            "1. Estructura y Layout:\n"
            "- Utiliza un contenedor principal que simule una hoja A4 con un ancho máximo de 850px y "
            "sombra suave (box-shadow) para visualización en pantalla.\n"
            "- Implementa un sistema de dos columnas mediante display:grid o flex: una columna izquierda "
            "principal (2/3 de ancho) para contenido detallado y una columna derecha (1/3 de ancho) tipo "
            "barra lateral con fondo gris muy suave (#f8fafc).\n"
            "- Crea un encabezado de ancho completo con fondo azul marino (#1e3a8a) y texto en blanco, centrado, "
            "que incluya el nombre en mayúsculas y datos de contacto.\n\n"
            "2. Estilo Visual y Tipografía:\n"
            "- Fuente: Sans-serif (Segoe UI o Arial), tamaño de cuerpo 11pt y un interlineado de 1.6 para máxima legibilidad.\n"
            "- Paleta de Colores: Azul profesional (#1e3a8a) para títulos y fondo de cabecera; gris pizarra para textos "
            "secundarios; y azul brillante (#3b82f6) para acentos.\n"
            "- Títulos: Las secciones (h2) deben estar en mayúsculas, con un borde inferior sólido del mismo color azul "
            "que la cabecera.\n\n"
            "3. Formato de Contenido:\n"
            "- Experiencia y Educación: Estructura cada ítem con el cargo/título en negrita, seguido de la "
            "empresa/institución en un color de acento, y las fechas en cursiva.\n"
            "- Listas: Las listas de la columna principal deben usar viñetas personalizadas (como ▹ o puntos azules). "
            "Las listas en la columna lateral deben aparecer como tags o etiquetas con fondo gris claro y bordes redondeados.\n"
            "- Resumen: Incluye una sección de Perfil Profesional con texto justificado.\n\n"
            "4. Adaptabilidad:\n"
            "- Asegúrate de que el CSS incluya una regla @media print para eliminar sombras, fondos de página y ajustar "
            "márgenes al exportar a PDF.\n\n"
            "Adapta estas reglas base al estilo elegido por el usuario sin perder profesionalismo.\n\n"
            "DATOS DEL CANDIDATO:\n"
            f"- Nombre completo: {full_name}\n"
            f"- Email: {email}\n"
            f"- Vacantes objetivo: {profile_data.get('target_roles', '')}\n"
            f"- Resumen profesional: {profile_data.get('summary', '')}\n"
            f"- Habilidades: {profile_data.get('skills', '')}\n"
            f"- Experiencia: {profile_data.get('experience', '')}\n"
            f"- Educación: {profile_data.get('education', '')}\n"
            f"- Idiomas: {profile_data.get('languages', '')}\n"
            f"- Certificaciones: {profile_data.get('certifications', '')}\n"
            f"- Logros: {profile_data.get('achievements', '')}\n\n"
            "PREFERENCIAS DE DISEÑO DINÁMICAS DEL USUARIO:\n"
            f"- Estilo seleccionado: {style_cfg['style_name']} ({selected_palette_id})\n"
            f"- Color principal de cabecera/títulos: {style_cfg['header_primary']}\n"
            f"- Color de acento secundario: {style_cfg['accent']}\n"
            f"- Color de texto principal: {style_cfg['text']}\n"
            f"- Color de texto secundario: {style_cfg['muted']}\n"
            f"- Fondo de barra lateral: {style_cfg['sidebar_bg']}\n"
            f"- Fondo de página: {style_cfg['page_bg']}\n"
            f"- Tamaño de fuente principal elegido: {font_size}\n"
            f"- Distribución solicitada: {layout_desc}\n"
            f"- Foto de perfil: {photo_instruction}\n"
            f"- Regla de layout adicional: {column_instruction}\n"
            f"- Compatibilidad de renderización: {pdf_compat_instruction}\n"
            f"- Formato de exportación del archivo final: {output_format.upper()}\n\n"
            "INSTRUCCIONES CRÍTICAS DE SALIDA:\n"
            "1. Genera un documento HTML5 válido y completo: incluye <!DOCTYPE html>, <html>, <head> y <body>.\n"
            "2. Coloca TODOS los estilos CSS dentro de un bloque <style> en el <head>. Sin archivos CSS externos.\n"
            "3. Usa fuente sans-serif (Arial, Helvetica o similar).\n"
            "4. Incluye TODAS las secciones en este orden: encabezado con nombre y contacto, perfil profesional, "
            "vacantes objetivo, experiencia laboral, educación, habilidades, idiomas, certificaciones y logros.\n"
            "5. Usa ÚNICAMENTE la información proporcionada en DATOS DEL CANDIDATO. No inventes datos.\n"
            "6. Aplica los colores, tamaño de fuente y layout especificados en PREFERENCIAS DE DISEÑO DINÁMICAS.\n"
            "7. Responde SOLO con el código HTML puro, sin explicaciones, sin markdown, "
            "sin texto previo ni posterior.\n"
            "El HTML generado debe ser directamente renderizable como un CV profesional real."
        )

        payload = {
            "contents": [{"parts": [{"text": html_prompt}]}],
            "generationConfig": {
                "temperature": 0.3,
                "topP": 0.85,
                "maxOutputTokens": 8192,
            },
        }
        models_to_try = [GEMINI_MODEL] + [m for m in GEMINI_FALLBACK_MODELS if m != GEMINI_MODEL]
        for model_name in models_to_try:
            url = (
                f"https://generativelanguage.googleapis.com/{GEMINI_API_VERSION}"
                f"/models/{model_name}:generateContent"
            )
            try:
                response = await asyncio.to_thread(
                    httpx.post,
                    url,
                    json=payload,
                    headers={"x-goog-api-key": GEMINI_API_KEY},
                    timeout=90.0,
                )
                response.raise_for_status()
                gemini_text = extract_gemini_text(response.json())
                if gemini_text:
                    return extract_cv_html_from_gemini_text(gemini_text)
            except httpx.HTTPStatusError as ex:
                if ex.response.status_code == 404:
                    continue
                detail = extract_gemini_error_detail(ex.response)
                raise ValueError(f"Gemini error {ex.response.status_code}: {detail or ''}") from None
            except httpx.HTTPError as ex:
                raise ValueError(f"No se pudo conectar con Gemini: {ex}") from None
        raise ValueError("Gemini no respondió en ningún modelo configurado.")

    async def finalize_flow():
        req_keys = [s["key"] for s in sections if s.get("required", False)]
        missing = [k for k in req_keys if not chat_state["profile_data"].get(k)]
        if missing:
            add_message("bot", f"Faltan datos obligatorios: {', '.join(missing)}.")
            page.update()
            return

        if not generate_and_store:
            chat_state["profile_data"]["output_format"] = chat_state["output_format"]
            await on_complete(chat_state["profile_data"])
            palette_label = get_ai_cv_color_palette_label(chat_state["profile_data"].get("cv_color_palette"))
            font_label = get_ai_cv_font_size_label(chat_state["profile_data"].get("cv_font_size"))
            columns_label = get_ai_cv_columns_label(chat_state["profile_data"].get("cv_columns"))
            photo_label = get_ai_cv_photo_option_label(chat_state["profile_data"].get("cv_include_photo"))
            add_message(
                "bot",
                "Listo. Ya guardé tu información para usarla al registrarte en "
                f"formato {chat_state['output_format'].upper()} con estilo: {palette_label}, {font_label}, {columns_label}, {photo_label}.",
            )
            status_text.value = "Información del CV preparada correctamente."
            page.update()
            return

        if not current_user:
            add_message("bot", "No hay usuario autenticado. Inicia sesión para generar y guardar tu CV.")
            page.update()
            return

        chat_state["processing"] = True
        set_choice_mode(False)
        set_input_enabled(False)
        add_message("bot", "Perfecto. Estoy generando tu CV con IA y guardándolo en Supabase...")
        status_text.value = "Generando CV con Gemini..."
        page.update()

        try:
            details = user_details.get(current_user or "", {})
            full_name = f"{details.get('nombre', '')} {details.get('apellido', '')}".strip()
            output_format = normalize_ai_cv_output_format(chat_state.get("output_format", "docx"))
            chat_state["profile_data"]["cv_color_palette"] = normalize_ai_cv_color_palette(
                chat_state["profile_data"].get("cv_color_palette", "")
            )
            chat_state["profile_data"]["cv_font_size"] = normalize_ai_cv_font_size(
                chat_state["profile_data"].get("cv_font_size", "")
            )
            chat_state["profile_data"]["cv_columns"] = normalize_ai_cv_columns(
                chat_state["profile_data"].get("cv_columns", "")
            )
            chat_state["profile_data"]["cv_include_photo"] = normalize_ai_cv_photo_option(
                chat_state["profile_data"].get("cv_include_photo", "")
            )
            resolved_name = full_name or current_user or "Usuario"

            status_text.value = "Solicitando HTML del CV a Gemini..."
            page.update()

            cv_html = await generate_cv_html_from_profile(
                chat_state["profile_data"],
                resolved_name,
                current_user or "",
                output_format,
            )

            print("\n===== AI CV BUILDER — GEMINI HTML START =====")
            print(f"Modelo: {GEMINI_MODEL} | Formato: {output_format.upper()}")
            print(cv_html)
            print("===== AI CV BUILDER — GEMINI HTML END =====\n")

            if not looks_like_cv_html(cv_html):
                raise ValueError("Gemini no devolvió un documento HTML de CV válido.")

            photo_bytes = extract_profile_photo_bytes(chat_state["profile_data"])
            if photo_bytes:
                cv_html = inject_photo_into_html(cv_html, photo_bytes)

            status_text.value = f"Convirtiendo HTML a {output_format.upper()}..."
            page.update()

            if output_format == "pdf":
                file_bytes = await asyncio.to_thread(create_cv_pdf_from_html, cv_html)
                file_ext = "pdf"
            else:
                file_bytes = await asyncio.to_thread(create_cv_docx_from_html, cv_html)
                file_ext = "docx"

            file_name = f"cv_{sanitize_storage_name(resolved_name)}.{file_ext}"
            profile_to_store = {k: v for k, v in chat_state["profile_data"].items() if k != "cv_photo_base64"}
            profile_to_store["output_format"] = output_format
            profile_to_store["cv_photo_uploaded"] = bool(photo_bytes)

            status_text.value = "Guardando CV en Supabase..."
            page.update()

            metadata = await asyncio.to_thread(
                upload_cv_bytes,
                current_user,
                file_name,
                file_bytes,
                "ai_generated",
                chat_state["profile_data"].get("target_roles", ""),
                profile_to_store,
            )
            metadata["output_format"] = output_format
            metadata["generated_code"] = cv_html[:500]
            await on_complete(metadata)

            public_url = metadata.get("public_url", "")
            add_message("bot", f"Tu CV fue generado y guardado: {file_name}")

            result_actions.controls.clear()
            if public_url:
                result_actions.controls.append(ft.TextButton("Abrir CV generado", url=public_url))
                add_message("bot", f"Puedes verlo aquí: {public_url}")
            else:
                add_message("bot", "El CV se guardó, pero no se pudo generar un enlace público.")

            status_text.value = f"CV {output_format.upper()} generado y almacenado en Supabase."
        except Exception as ex:
            add_message("bot", f"No pude generar o guardar el CV con IA: {ex}")
            status_text.value = "Error al generar/guardar el CV."
            set_input_enabled(True)
            chat_state["processing"] = False

        page.update()

    async def submit_answer(e):
        if chat_state["processing"]:
            return

        if chat_state.get("awaiting_choice"):
            return

        user_input = (input_field.value or "").strip()
        if not user_input:
            return

        input_field.value = ""

        current_section_idx = chat_state["current_section"]
        if current_section_idx >= len(sections):
            return

        section = sections[current_section_idx]
        current_key = section["key"]
        min_length = section.get("min_length", 5)

        add_message("user", user_input)

        if len(user_input) < min_length:
            add_message("bot", f"Tu respuesta es un poco corta. Por favor, agrega más detalles (al menos {min_length} caracteres).")
            page.update()
            return

        chat_state["profile_data"][current_key] = user_input

        chat_state["current_section"] += 1
        if chat_state["current_section"] < len(sections):
            ask_current_question()
        else:
            ask_next_design_choice()
        page.update()

    def on_photo_picker_result(evt):
        if chat_state.get("awaiting_choice") != "photo_upload":
            return

        if not evt.files:
            add_message("bot", "No se seleccionó ninguna imagen.")
            page.update()
            return

        selected = evt.files[0]
        selected_path = selected.path or ""
        selected_name = selected.name or Path(selected_path).name

        if not selected_path or not os.path.isfile(selected_path):
            add_message("bot", "No se pudo leer el archivo seleccionado.")
            page.update()
            return

        if not is_supported_photo_file(selected_name):
            add_message("bot", "Formato no soportado. Usa PNG, JPG o JPEG.")
            page.update()
            return

        try:
            with open(selected_path, "rb") as image_file:
                raw_bytes = image_file.read()
        except Exception:
            add_message("bot", "No se pudo cargar la foto desde el archivo.")
            page.update()
            return

        if not raw_bytes:
            add_message("bot", "La imagen seleccionada está vacía.")
            page.update()
            return

        if len(raw_bytes) > MAX_PROFILE_PHOTO_BYTES:
            add_message("bot", "La foto supera 3 MB. Elige una imagen más liviana.")
            page.update()
            return

        chat_state["profile_data"]["cv_photo_base64"] = base64.b64encode(raw_bytes).decode("ascii")
        add_message("user", f"Foto cargada: {selected_name}")
        add_message("bot", "Perfecto, ya tengo tu foto. Pulsa Continuar para seguir.")
        page.update()

    photo_picker.on_result = on_photo_picker_result

    # Crear diálogo
    dlg = ft.AlertDialog(
        title=ft.Text(" "),
        content=ft.Column([
            ft.Text(
                "Responde paso a paso. Enter para enviar.",
                color=palette["muted"],
                size=11,
                weight=ft.FontWeight.W_500,
            ),
            ft.Divider(height=1, color=palette["border"]),
            ft.Container(
                expand=True,
                content=messages_list,
            ),
            result_actions,
            choice_quick_actions,
            status_text,
            ft.Container(
                content=ft.Row([
                    input_field,
                    send_button,
                ], spacing=8, expand=True, vertical_alignment=ft.CrossAxisAlignment.END),
            ),
        ], spacing=8, expand=True),
        actions=[
            ft.TextButton("Cancelar", on_click=lambda e: close_dialog(e)),
        ],
        modal=False,
        inset_padding=ft.Padding(top=20, left=20, right=20, bottom=20),
    )

    send_button.on_click = lambda e: run_async(submit_answer(e))
    # Configurar submit en Enter
    input_field.on_submit = lambda e: run_async(submit_answer(e))

    open_dialog(dlg)

    add_message("bot", "Hola. Vamos a construir tu CV paso a paso.")
    ask_current_question()
    
    page.update()
    
    # Nota: No intentamos hacer focus() en controles dentro de AlertDialog
    # ya que causa RuntimeError hasta que se renderice completamente.
    # El TextField capturará Enter key input automáticamente.


def open_gemini_chat_cv_dialog(
    page: ft.Page,
    palette: dict[str, str],
    on_complete,
    generate_and_store: bool = False,
):
    """Diálogo de chat completamente guiado por Gemini para construir un CV de forma conversacional."""
    dlg: ft.AlertDialog | None = None
    chat_history: list[dict] = []
    chat_state = {
        "processing": False,
        "output_format": "docx",
        "awaiting_format": False,
        "ready": False,
        "startup_loading": True,
    }

    current_details = user_details.get(current_user or "", {})
    prefilled_name = normalize_text(
        f"{current_details.get('nombre', '')} {current_details.get('apellido', '')}".strip()
    )
    prefilled_email = normalize_text(current_user or "")
    prefilled_phone = normalize_text(
        current_details.get("telefono", "")
        or current_details.get("phone", "")
        or current_details.get("telefono_movil", "")
    )

    INITIAL_PROMPT = (
        "Actúa como un experto en reclutamiento técnico. Ayúdame a construir mi CV sección por sección "
        "siguiendo estas reglas estrictas:\n\n"
        "Brevedad: No hagas introducciones largas ni me des la bienvenida extendida. Ve directo a la acción.\n\n"
        "Micro-pasos: Solicítame la información de una sola sección a la vez. "
        "No pases a la siguiente hasta que hayamos redactado y aprobado la anterior.\n\n"
        "Optimización: Redacta mis logros usando la fórmula de Google (Acción + Medida + Resultado).\n\n"
        "Diseño: Al final, ofréceme un abanico de 3 estilos visuales con paletas de colores.\n\n"
        "Empecemos ahora: Solicítame únicamente los datos necesarios para la Sección 1 "
        "(Datos personales y Título profesional) de forma breve."
    )

    messages_list = ft.ListView(spacing=8, expand=True, auto_scroll=True)
    startup_loading_overlay = mk_loading_overlay(
        page,
        palette,
        "Iniciando conversación con Gemini",
        "Estamos preparando la primera pregunta para construir tu CV.",
        visible=True,
    )
    result_actions = ft.Row([], wrap=True, spacing=8)
    status_text = mk_status_text(page, palette)
    format_quick_actions = ft.Row([], wrap=True, spacing=8, visible=False)
    generate_button = mk_button(
        "Generar CV",
        variant="primary",
        palette=palette,
    )
    generate_button.visible = False

    input_field = mk_textfield(
        "Tu mensaje (Enter = nueva línea · botón ➤ para enviar)",
        palette=palette,
        multiline=True,
        min_lines=1,
        max_lines=5,
    )
    send_button = ft.IconButton(
        icon=ft.Icons.SEND,
        icon_color=palette["accent"],
        tooltip="Enviar mensaje",
    )

    def run_async(coro):
        task = asyncio.create_task(coro)

        def _consume(done_task: asyncio.Task):
            try:
                err = done_task.exception()
                if err and "Control must be added" not in str(err):
                    pass
            except Exception:
                pass

        task.add_done_callback(_consume)

    def open_dialog(dialog: ft.AlertDialog):
        if dialog not in page.overlay:
            page.overlay.append(dialog)
        dialog.open = True
        page.update()

    def close_dialog(e=None):
        nonlocal dlg
        if dlg is None:
            return
        dlg.open = False
        page.update()

    def build_message_text(content: str, color: str, align: ft.TextAlign) -> ft.Text:
        normalized_content = content or ""
        parts = re.split(r"(\*\*[^*]+\*\*)", normalized_content)
        spans: list[ft.TextSpan] = []

        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                spans.append(
                    ft.TextSpan(
                        text=part[2:-2],
                        style=ft.TextStyle(color=color, size=12, weight=ft.FontWeight.W_700),
                    )
                )
            else:
                spans.append(
                    ft.TextSpan(
                        text=part,
                        style=ft.TextStyle(color=color, size=12),
                    )
                )

        if spans:
            return ft.Text(spans=spans, selectable=True, text_align=align)

        return ft.Text(
            normalized_content.replace("**", ""),
            color=color,
            size=12,
            selectable=True,
            text_align=align,
        )

    def add_message(role: str, content: str):
        if role == "bot":
            bubble = mk_chat_bubble(
                page,
                palette,
                content=build_message_text(content, palette["bg"], ft.TextAlign.LEFT),
                role="bot",
            )
        else:
            bubble = mk_chat_bubble(
                page,
                palette,
                content=build_message_text(content, palette["text"], ft.TextAlign.RIGHT),
                role="user",
            )
        messages_list.controls.append(bubble)

    def set_input_enabled(enabled: bool):
        input_field.disabled = not enabled
        send_button.disabled = not enabled

    async def call_gemini(user_text: str) -> str:
        chat_history.append({"role": "user", "parts": [{"text": user_text}]})
        payload = {
            "contents": chat_history,
            "generationConfig": {
                "temperature": 0.7,
                "topP": 0.9,
                "maxOutputTokens": 2048,
            },
        }
        models_to_try = [GEMINI_MODEL] + [m for m in GEMINI_FALLBACK_MODELS if m != GEMINI_MODEL]
        for model_name in models_to_try:
            url = (
                f"https://generativelanguage.googleapis.com/{GEMINI_API_VERSION}"
                f"/models/{model_name}:generateContent"
            )
            try:
                response = await asyncio.to_thread(
                    httpx.post,
                    url,
                    json=payload,
                    headers={"x-goog-api-key": GEMINI_API_KEY},
                    timeout=60.0,
                )
                response.raise_for_status()
                gemini_text = extract_gemini_text(response.json())
                if gemini_text:
                    chat_history.append({"role": "model", "parts": [{"text": gemini_text}]})
                    return gemini_text
            except httpx.HTTPStatusError as ex:
                if ex.response.status_code == 404:
                    continue
                detail = extract_gemini_error_detail(ex.response)
                raise ValueError(f"Gemini error {ex.response.status_code}: {detail or ''}") from None
            except httpx.HTTPError as ex:
                raise ValueError(f"No se pudo conectar con Gemini: {ex}") from None
        raise ValueError("Gemini no respondió en ningún modelo configurado.")

    async def extract_cv_html_from_chat() -> str:
        """Envía un prompt final a Gemini solicitando el HTML completo del CV basado en la conversación."""
        selected_output_format = normalize_ai_cv_output_format(chat_state.get("output_format", "docx"))

        # Pistas visuales acordadas en la conversación; si el usuario las escogió mediante botones
        # antes de abrir este diálogo estarán en chat_state, de lo contrario Gemini las decide.
        palette_css_map = {
            "azul_profesional": {"accent": "#1F4E79", "text": "#111827", "bg": "#EFF6FF"},
            "verde_moderno":    {"accent": "#166534", "text": "#0F172A", "bg": "#F0FDF4"},
            "gris_ejecutivo":   {"accent": "#374151", "text": "#111827", "bg": "#F9FAFB"},
        }
        font_pt_map = {"compacta": "10pt", "estandar": "11pt", "amplia": "13pt"}
        columns_desc_map = {"una_columna": "una sola columna", "dos_columnas": "dos columnas lado a lado"}

        selected_palette_id = normalize_ai_cv_color_palette(chat_state.get("cv_color_palette", ""))
        selected_font_id    = normalize_ai_cv_font_size(chat_state.get("cv_font_size", ""))
        selected_columns_id = normalize_ai_cv_columns(chat_state.get("cv_columns", ""))
        include_photo       = normalize_ai_cv_photo_option(chat_state.get("cv_include_photo", "")) == "con_foto"

        colors      = palette_css_map.get(selected_palette_id, palette_css_map["azul_profesional"])
        font_size   = font_pt_map.get(selected_font_id, "11pt")
        layout_desc = columns_desc_map.get(selected_columns_id, "una sola columna")

        photo_instruction = (
            'Incluye un elemento <img id="cv-photo" alt="Foto de perfil" '
            'style="width:110px;height:130px;object-fit:cover;border-radius:4px;'
            'float:right;margin:0 0 12px 16px;" />'
        ) if include_photo else "No incluyas elemento de foto."

        extraction_prompt = (
            "Objetivo: generar el código HTML5 completo de un CV profesional personalizado para el usuario, "
            "usando TODA la información compartida en esta conversación "
            "(datos personales, experiencia, educación, habilidades, idiomas, certificaciones, logros y vacantes objetivo).\n\n"
            "Instrucciones críticas:\n"
            "1. Genera un documento HTML5 válido y completo: incluye <!DOCTYPE html>, <html>, <head> y <body>.\n"
            "2. Coloca TODOS los estilos CSS dentro de un bloque <style> en el <head>. Sin CSS externo.\n"
            f"3. Color de acento para encabezados y títulos de sección: {colors['accent']}.\n"
            f"4. Color del cuerpo del texto: {colors['text']}.\n"
            f"5. Color de fondo del documento: {colors['bg']}.\n"
            f"6. Tamaño de fuente principal: {font_size}. Usa fuente sans-serif (Arial, Helvetica o similar).\n"
            f"7. Diseño de página: {layout_desc}.\n"
            f"8. Foto de perfil: {photo_instruction}\n"
            "9. Incluye TODAS las secciones: nombre completo y datos de contacto en el encabezado, "
            "luego perfil profesional, vacantes objetivo, experiencia laboral, educación, "
            "habilidades, idiomas, certificaciones y logros.\n"
            "10. Usa ÚNICAMENTE la información que el usuario proporcionó; no inventes datos.\n"
            f"11. El documento será exportado como {selected_output_format.upper()}.\n"
            "12. Responde SOLO con el código HTML, sin explicaciones, sin markdown, sin texto previo ni posterior.\n"
            "El HTML que generes debe ser directamente renderizable y representar un CV profesional real.\n"
        )

        temp_history = chat_history + [{"role": "user", "parts": [{"text": extraction_prompt}]}]
        payload = {
            "contents": temp_history,
            "generationConfig": {
                "temperature": 0.2,
                "topP": 0.85,
                "maxOutputTokens": 8192,
            },
        }
        models_to_try = [GEMINI_MODEL] + [m for m in GEMINI_FALLBACK_MODELS if m != GEMINI_MODEL]
        for model_name in models_to_try:
            url = (
                f"https://generativelanguage.googleapis.com/{GEMINI_API_VERSION}"
                f"/models/{model_name}:generateContent"
            )
            try:
                response = await asyncio.to_thread(
                    httpx.post,
                    url,
                    json=payload,
                    headers={"x-goog-api-key": GEMINI_API_KEY},
                    timeout=90.0,
                )
                response.raise_for_status()
                gemini_text = extract_gemini_text(response.json())
                if gemini_text:
                    cv_html = extract_cv_html_from_gemini_text(gemini_text)
                    print("\n===== GEMINI CV HTML START =====")
                    print(f"Modelo: {model_name}")
                    print(cv_html)
                    print("===== GEMINI CV HTML END =====\n")
                    return cv_html
            except httpx.HTTPStatusError as ex:
                if ex.response.status_code == 404:
                    continue
                raise
        raise ValueError("No se pudo generar el HTML del CV desde el historial de chat.")

    async def finalize_generation():
        if generate_and_store and not current_user:
            add_message("bot", "No hay usuario autenticado. Inicia sesión primero.")
            page.update()
            return

        chat_state["processing"] = True
        set_input_enabled(False)
        status_text.value = "Generando el HTML del CV con Gemini..."
        page.update()

        try:
            cv_html = await extract_cv_html_from_chat()
            if not looks_like_cv_html(cv_html):
                raise ValueError("Gemini no devolvió un documento HTML de CV válido.")

            # Inyectar foto real si el usuario la subió durante la conversación.
            photo_bytes = extract_profile_photo_bytes(chat_state.get("profile_data") or {})
            if photo_bytes:
                cv_html = inject_photo_into_html(cv_html, photo_bytes)

            output_format = normalize_ai_cv_output_format(chat_state.get("output_format", "docx"))
            status_text.value = f"Convirtiendo HTML a {output_format.upper()}..."
            page.update()

            if output_format == "pdf":
                file_bytes = await asyncio.to_thread(create_cv_pdf_from_html, cv_html)
                file_ext = "pdf"
            else:
                file_bytes = await asyncio.to_thread(create_cv_docx_from_html, cv_html)
                file_ext = "docx"

            if generate_and_store:
                details = user_details.get(current_user or "", {})
                resolved_name = (
                    f"{details.get('nombre', '')} {details.get('apellido', '')}".strip()
                    or current_user
                    or "Usuario"
                )
                resolved_email = current_user or ""

                file_name = f"cv_ia_{sanitize_storage_name(resolved_name or resolved_email)}.{file_ext}"
                metadata = await asyncio.to_thread(
                    upload_cv_bytes,
                    resolved_email,
                    file_name,
                    file_bytes,
                    "ai_chat_generated",
                    "",
                    {"output_format": output_format},
                )
                metadata["output_format"] = output_format
                metadata["generated_code"] = cv_html[:500]  # prefijo para trazabilidad
                await on_complete(metadata)

                public_url = metadata.get("public_url", "")
                add_message("bot", f"Tu CV fue generado y guardado: {file_name}")
                result_actions.controls.clear()
                if public_url:
                    result_actions.controls.append(ft.TextButton("Abrir CV generado", url=public_url))
                    add_message("bot", f"Enlace público: {public_url}")
                status_text.value = f"CV {output_format.upper()} generado y almacenado."
            else:
                await on_complete({"output_format": output_format, "cv_html": cv_html})
                add_message("bot", "CV en HTML preparado para el registro.")
                status_text.value = "CV listo."

        except Exception as ex:
            add_message("bot", f"Error al generar el CV: {ex}")
            status_text.value = "Error al generar el CV."
            generate_button.visible = True

        chat_state["processing"] = False
        set_input_enabled(True)
        page.update()

    async def choose_format(fmt: str):
        if chat_state["processing"] or not chat_state["awaiting_format"]:
            return
        chat_state["output_format"] = fmt
        chat_state["awaiting_format"] = False
        format_quick_actions.visible = False
        add_message("user", fmt.upper())
        add_message("bot", f"Formato {fmt.upper()} seleccionado. Generando tu CV...")
        page.update()
        await finalize_generation()

    docx_fmt_btn = mk_button(
        "Docx",
        on_click=lambda e: run_async(choose_format("docx")),
        variant="secondary",
        palette=palette,
    )
    pdf_fmt_btn = mk_button(
        "Pdf",
        on_click=lambda e: run_async(choose_format("pdf")),
        variant="secondary",
        palette=palette,
    )
    format_quick_actions.controls = [docx_fmt_btn, pdf_fmt_btn]

    async def on_generate_cv(e):
        if chat_state["processing"]:
            return
        generate_button.visible = False
        chat_state["awaiting_format"] = True
        format_quick_actions.visible = True
        add_message(
            "bot",
            "¿En qué formato deseas tu CV?\n1) DOCX (editable)\n2) PDF\n\nUsa los botones o escribe: docx o pdf",
        )
        page.update()

    generate_button.on_click = lambda e: run_async(on_generate_cv(e))

    async def handle_send(e):
        if chat_state["processing"]:
            return

        user_text = (input_field.value or "").strip()
        if not user_text:
            return
        input_field.value = ""

        if chat_state["awaiting_format"]:
            add_message("user", user_text)
            alias_map = {"1": "docx", "docx": "docx", "2": "pdf", "pdf": "pdf"}
            selected = alias_map.get(user_text.lower())
            if not selected:
                add_message("bot", "Por favor responde con: docx o pdf (o usa los botones).")
                page.update()
                return
            chat_state["awaiting_format"] = False
            format_quick_actions.visible = False
            add_message("bot", f"Formato {selected.upper()} seleccionado. Generando tu CV...")
            chat_state["output_format"] = selected
            page.update()
            await finalize_generation()
            return

        add_message("user", user_text)
        chat_state["processing"] = True
        set_input_enabled(False)
        status_text.value = "Gemini está escribiendo..."
        page.update()

        try:
            response_text = await call_gemini(user_text)
            add_message("bot", response_text)
            if not chat_state["ready"]:
                chat_state["ready"] = True
                generate_button.visible = True
        except Exception as ex:
            add_message("bot", f"Error al contactar Gemini: {ex}")

        chat_state["processing"] = False
        status_text.value = ""
        set_input_enabled(True)
        page.update()

    send_button.on_click = lambda e: run_async(handle_send(e))

    dlg = ft.AlertDialog(
        title=ft.Text(" "),
        content=ft.Column(
            [
                ft.Text(
                    "Crear CV con IA — conversa con Gemini para construir tu CV.",
                    color=palette["muted"],
                    size=11,
                    weight=ft.FontWeight.W_500,
                ),
                ft.Divider(height=1, color=palette["border"]),
                ft.Stack(
                    [
                        ft.Container(expand=True, content=messages_list),
                        startup_loading_overlay,
                    ],
                    expand=True,
                ),
                result_actions,
                format_quick_actions,
                generate_button,
                status_text,
                ft.Container(
                    content=ft.Row(
                        [input_field, send_button],
                        spacing=8,
                        expand=True,
                        vertical_alignment=ft.CrossAxisAlignment.END,
                    ),
                ),
            ],
            spacing=8,
            expand=True,
        ),
        actions=[ft.TextButton("Cancelar", on_click=lambda e: close_dialog(e))],
        modal=False,
        inset_padding=ft.Padding(top=20, left=20, right=20, bottom=20),
    )

    open_dialog(dlg)

    async def start_conversation():
        chat_state["processing"] = True
        set_input_enabled(False)
        status_text.value = "Iniciando conversación con Gemini..."
        startup_loading_overlay.visible = True
        page.update()
        try:
            if not GEMINI_API_KEY:
                raise ValueError("GEMINI_API_KEY no está configurada.")
            response_text = await call_gemini(INITIAL_PROMPT)
            add_message("bot", response_text)
            chat_state["ready"] = True
            generate_button.visible = True
        except Exception as ex:
            add_message("bot", f"No se pudo iniciar la conversación con Gemini: {ex}")
        startup_loading_overlay.visible = False
        chat_state["startup_loading"] = False
        chat_state["processing"] = False
        status_text.value = ""
        set_input_enabled(True)
        page.update()

    run_async(start_conversation())
    page.update()


def set_authenticated_bottom_appbar(page: ft.Page, palette: dict[str, str], active_section: str | None = None):
    def make_nav_button(icon, tooltip, on_click, section_key=None, danger=False):
        is_active = section_key is not None and active_section == section_key
        return ft.IconButton(
            icon=icon,
            tooltip=tooltip,
            on_click=lambda e: asyncio.create_task(on_click()),
            icon_color=palette["accent_text"] if is_active else (palette["danger"] if danger else palette["text"]),
            style=get_nav_icon_button_style(is_active, palette),
        )

    async def go_search():
        await show_main_ui(page)

    async def go_add():
        await show_add_manual_job_dialog(page)

    async def go_cv():
        await show_cv_prep_ui(page)

    async def do_logout():
        await logout_user(page)

    page.bottom_appbar = ft.BottomAppBar(
        bgcolor=palette["surface"],
        elevation=8,
        padding=ft.Padding(12, 8, 12, 8),
        content=ft.Row(
            [
                make_nav_button(ft.Icons.SEARCH, "Búsqueda de empleo", go_search, section_key="search"),
                make_nav_button(ft.Icons.ADD_CIRCLE_OUTLINE, "Agregar empleo manualmente", go_add),
                make_nav_button(ft.Icons.DESCRIPTION_OUTLINED, "Mi Curriculum", go_cv, section_key="cv"),
                make_nav_button(ft.Icons.LOGOUT, "Cerrar sesión", do_logout, danger=True),
            ],
            alignment=ft.MainAxisAlignment.SPACE_AROUND,
        ),
    )


async def show_cv_prep_ui(page: ft.Page):
    page.controls.clear()
    page.overlay.clear()
    page.title = "JobFriends - Preparar CV"
    page.scroll = ft.ScrollMode.ADAPTIVE
    mobile = is_mobile(page)
    tablet = is_tablet(page)
    page.padding = get_page_padding(page)
    if page.theme_mode == ft.ThemeMode.SYSTEM:
        page.theme_mode = ft.ThemeMode.LIGHT
    palette = get_palette(page)
    page.bgcolor = palette["bg"]

    async def toggle_theme(e):
        page.theme_mode = ft.ThemeMode.LIGHT if is_dark_mode(page) else ft.ThemeMode.DARK
        await show_cv_prep_ui(page)

    # Evita reconstrucciones completas de la vista durante interacciones del dialogo.
    page.on_resized = lambda e: page.update()
    set_authenticated_bottom_appbar(page, palette, active_section="cv")

    cv_record = get_applicant_cv(current_user) if current_user else None
    cv_path_input = mk_textfield(
        "Ruta local del CV (.pdf, .doc, .docx)",
        palette=palette,
    )
    cv_status = mk_status_text(page, palette, "Sube un CV en PDF/DOC/DOCX o genera una versión optimizada con IA.")
    cv_info = ft.Column(spacing=8)

    def refresh_cv_info(record: dict | None = None):
        current_record = record if record is not None else cv_record
        cv_info.controls.clear()
        if current_record:
            source_label = "IA" if current_record.get("source") == "ai_generated" else "Archivo cargado"
            cv_info.controls.extend([
                ft.Text(f"CV actual: {current_record.get('file_name', 'Sin nombre')}", color=palette["text"], weight=ft.FontWeight.W_700),
                ft.Text(f"Origen: {source_label}", color=palette["muted"]),
                ft.Text(f"Vacantes objetivo: {current_record.get('target_roles') or 'No especificadas'}", color=palette["muted"]),
            ])
            public_url = current_record.get("public_url", "")
            if public_url:
                cv_info.controls.append(ft.TextButton("Abrir CV guardado", url=public_url))
        else:
            cv_info.controls.append(
                mk_empty_state(
                    page,
                    palette,
                    "Todavía no tienes un CV guardado.",
                    "Carga un archivo existente o genera uno nuevo con IA para empezar.",
                )
            )

    async def save_uploaded_cv():
        nonlocal cv_record
        file_path = (cv_path_input.value or "").strip().strip('"')
        if not file_path:
            cv_status.value = "Ingresa la ruta local del archivo PDF, DOC o DOCX."
            cv_status.color = palette["danger"]
            page.update()
            return

        if not os.path.isfile(file_path):
            cv_status.value = "La ruta indicada no existe o no es un archivo válido."
            cv_status.color = palette["danger"]
            page.update()
            return

        file_name = Path(file_path).name
        if not is_supported_cv_file(file_name):
            cv_status.value = "Solo se admiten archivos PDF, DOC o DOCX."
            cv_status.color = palette["danger"]
            page.update()
            return

        try:
            metadata = await asyncio.to_thread(upload_cv_file, current_user, file_path, file_name)
            cv_record = metadata
            cv_status.value = "CV cargado y guardado correctamente."
            cv_status.color = palette["success"]
            refresh_cv_info(cv_record)
        except Exception:
            cv_status.value = "No se pudo guardar el CV en Supabase."
            cv_status.color = palette["danger"]
        page.update()

    async def save_ai_cv(metadata: dict):
        nonlocal cv_record
        try:
            cv_record = metadata
            cv_status.value = "Borrador de CV generado y guardado correctamente."
            cv_status.color = palette["success"]
            refresh_cv_info(cv_record)
        except Exception:
            cv_status.value = "No se pudo generar o guardar el CV con IA."
            cv_status.color = palette["danger"]
        page.update()

    def open_ai_cv_chat(e):
        try:
            open_ai_cv_builder_dialog(
                page,
                palette,
                (cv_record or {}).get("profile_data") or {},
                save_ai_cv,
                generate_and_store=True,
            )
        except Exception:
            cv_status.value = "No se pudo abrir el chat de CV con IA. Reintenta."
            cv_status.color = palette["danger"]
            page.update()

    def open_gemini_cv_chat(e):
        try:
            open_gemini_chat_cv_dialog(
                page,
                palette,
                save_ai_cv,
                generate_and_store=True,
            )
        except Exception:
            cv_status.value = "No se pudo abrir el chat guiado por Gemini. Reintenta."
            cv_status.color = palette["danger"]
            page.update()

    refresh_cv_info(cv_record)

    page.add(
        ft.Column([
            mk_surface_card(
                page,
                palette,
                variant="hero",
                content=ft.Row([
                    mk_section_header(page, palette, "Preparar CV", "Crea, actualiza o reemplaza tu currículum desde un solo lugar."),
                    make_theme_toggle(page, toggle_theme),
                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
            ),
            mk_surface_card(
                page,
                palette,
                variant="secondary",
                content=ft.Column([
                    mk_section_header(page, palette, "Gestiona tu curriculum", "Sube un CV existente o genera una versión optimizada con IA."),
                    ft.Text(
                        "Puedes subir un CV existente o crear un borrador con IA. Si Gemini está configurado, la generación se hace con el modelo real.",
                        size=ui_font_size(page, "body"),
                        color=palette["muted"],
                    ),
                    cv_info,
                    ft.ResponsiveRow([
                        ft.Container(
                            col={"xs": 12, "md": 12},
                            content=cv_path_input,
                        ),
                        ft.Container(
                            col={"xs": 12, "md": 4},
                            content=mk_button(
                                "Guardar CV cargado",
                                on_click=lambda e: asyncio.create_task(save_uploaded_cv()),
                                variant="primary",
                                palette=palette,
                            ),
                        ),
                        ft.Container(
                            col={"xs": 12, "md": 4},
                            content=mk_button(
                                "Crear con IA",
                                on_click=open_ai_cv_chat,
                                variant="secondary",
                                palette=palette,
                            ),
                        ),
                        ft.Container(
                            col={"xs": 12, "md": 4},
                            content=mk_button(
                                "Crear CV con IA",
                                on_click=open_gemini_cv_chat,
                                variant="primary",
                                palette=palette,
                            ),
                        ),
                    ], run_spacing=10),
                    ft.Text(
                        "La generación usa Gemini cuando GEMINI_API_KEY está configurada. Si falla, la app usa un borrador local como respaldo.",
                        size=12,
                        color=palette["muted"],
                    ),
                    cv_status,
                ], spacing=12),
            ),
        ], expand=True, spacing=12)
    )


async def show_interview_prep_ui(page: ft.Page):
    global pending_interview_application_id

    page.controls.clear()
    page.overlay.clear()
    page.title = "JobFriends - Preparar Entrevista"
    page.scroll = ft.ScrollMode.ADAPTIVE
    mobile = is_mobile(page)
    tablet = is_tablet(page)
    page.padding = get_page_padding(page)
    if page.theme_mode == ft.ThemeMode.SYSTEM:
        page.theme_mode = ft.ThemeMode.LIGHT
    palette = get_palette(page)
    page.bgcolor = palette["bg"]

    async def toggle_theme(e):
        page.theme_mode = ft.ThemeMode.LIGHT if is_dark_mode(page) else ft.ThemeMode.DARK
        await show_interview_prep_ui(page)

    page.on_resized = lambda e: page.update()
    set_authenticated_bottom_appbar(page, palette)

    if not current_user:
        page.add(
            mk_surface_card(
                page,
                palette,
                content=mk_empty_state(
                    page,
                    palette,
                    "Necesitas iniciar sesión para preparar entrevistas.",
                    "Accede con tu cuenta para abrir tus aplicaciones y practicar por vacante.",
                ),
            )
        )
        return

    try:
        response = supabase.table("Applications").select("*").eq("applicant_email", current_user).execute()
        applications = response.data or []
    except Exception:
        applications = []

    if not applications:
        page.add(
            ft.Column([
                mk_surface_card(
                    page,
                    palette,
                    variant="hero",
                    content=ft.Row([
                        mk_section_header(page, palette, "Preparar Entrevista", "Practica con contexto real de tus vacantes y mejora tus respuestas."),
                        make_theme_toggle(page, toggle_theme),
                    ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
                ),
                mk_surface_card(
                    page,
                    palette,
                    variant="secondary",
                    content=ft.Column([
                        mk_empty_state(
                            page,
                            palette,
                            "Aún no tienes empleos aplicados.",
                            "Registra una aplicación desde la búsqueda para empezar a practicar entrevistas.",
                        ),
                        mk_button("Volver al menú", on_click=lambda e: asyncio.create_task(show_menu_ui(page)), variant="secondary", palette=palette),
                    ], spacing=10),
                ),
            ], spacing=12)
        )
        return

    application_map = {str(item.get("id")): item for item in applications if item.get("id") is not None}
    application_ids = list(application_map.keys())
    preferred_id = None

    if pending_interview_application_id and pending_interview_application_id in application_map:
        preferred_id = pending_interview_application_id
    else:
        for app_id, app_row in application_map.items():
            if normalize_text(app_row.get("status", "")) == "En proceso":
                preferred_id = app_id
                break
    if not preferred_id:
        preferred_id = application_ids[0]

    pending_interview_application_id = None
    selected_application_id = preferred_id

    messages_column = ft.Column(spacing=8, scroll=ft.ScrollMode.ADAPTIVE, expand=True)
    status_text = mk_status_text(page, palette)
    description_text = ft.Text("", color=palette["text"], selectable=True)
    job_title_text = ft.Text("", size=18, weight=ft.FontWeight.W_700, color=palette["text"])
    job_meta_text = ft.Text("", size=12, color=palette["muted"])
    message_input = mk_textfield(
        "Escribe tu pregunta de entrevista",
        palette=palette,
        expand=True,
    )

    def get_selected_application() -> dict:
        return application_map.get(selected_application_id, applications[0])

    def get_chat_history(app_id: str) -> list[dict[str, str]]:
        history = interview_chat_state_by_application.setdefault(app_id, [])
        if not history:
            app_row = application_map.get(app_id, {})
            history.append({
                "role": "assistant",
                "content": (
                    f"¡Hola! Te ayudaré a preparar la entrevista para el cargo '{get_application_title(app_row)}'. "
                    "Cuéntame qué quieres practicar primero: presentación, preguntas técnicas o preguntas de comportamiento."
                ),
            })
        return history

    def render_chat(app_id: str):
        history = get_chat_history(app_id)
        messages_column.controls.clear()
        for message in history:
            is_user = message.get("role") == "user"
            align = ft.MainAxisAlignment.END if is_user else ft.MainAxisAlignment.START
            text_color = palette["text"] if is_user else palette["bg"]
            messages_column.controls.append(
                ft.Row([
                    mk_chat_bubble(
                        page,
                        palette,
                        content=ft.Text(message.get("content", ""), color=text_color, selectable=True),
                        width=480 if not mobile else None,
                        role="user" if is_user else "bot",
                    )
                ], alignment=align)
            )

    def refresh_selected_application_ui():
        app_row = get_selected_application()
        job_title_text.value = get_application_title(app_row)
        app_status = normalize_text(app_row.get("status", "")) or "Aplicado"
        app_link = normalize_text(app_row.get("application_link", ""))
        job_meta_text.value = f"Estado: {app_status} | Enlace: {app_link or 'No disponible'}"
        description_text.value = get_application_description(app_row) or "Sin descripción enriquecida guardada."
        render_chat(selected_application_id)

    async def on_application_change(e):
        nonlocal selected_application_id
        new_id = str(e.control.value or "")
        if new_id in application_map:
            selected_application_id = new_id
            refresh_selected_application_ui()
            page.update()

    async def send_message(e):
        user_message = normalize_text(message_input.value or "")
        if not user_message:
            return

        app_row = get_selected_application()
        history = get_chat_history(selected_application_id)
        history.append({"role": "user", "content": user_message})
        message_input.value = ""
        status_text.value = "Gemini está preparando tu respuesta..."
        render_chat(selected_application_id)
        page.update()

        try:
            reply_text = await asyncio.to_thread(
                generate_interview_reply_with_gemini,
                app_row,
                history,
                user_message,
            )
            history.append({"role": "assistant", "content": normalize_text(reply_text)})
            status_text.value = ""
        except Exception as ex:
            history.append({
                "role": "assistant",
                "content": "No pude responder con Gemini en este momento. Intenta de nuevo en unos segundos.",
            })
            status_text.value = f"Error: {ex}"

        render_chat(selected_application_id)
        page.update()

    async def go_back(e):
        await show_menu_ui(page)

    application_selector = ft.Dropdown(
        value=selected_application_id,
        options=[
            ft.dropdown.Option(
                key=str(app.get("id")),
                text=f"{get_application_title(app)} ({normalize_text(app.get('status', '')) or 'Aplicado'})",
            )
            for app in applications
            if app.get("id") is not None
        ],
        width=520 if not mobile else None,
    )
    application_selector.on_change = on_application_change

    refresh_selected_application_ui()

    page.add(
        ft.Column([
            mk_surface_card(
                page,
                palette,
                variant="hero",
                content=ft.Row([
                    mk_section_header(page, palette, "Preparar Entrevista", "Selecciona una vacante, revisa el contexto y ensaya por chat."),
                    make_theme_toggle(page, toggle_theme),
                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
            ),
            mk_surface_card(
                page,
                palette,
                variant="secondary",
                content=ft.Column([
                    mk_section_header(page, palette, "Selecciona la oferta para practicar", "Consulta la descripción enriquecida y entrena tu conversación antes de la entrevista."),
                    application_selector,
                    job_title_text,
                    job_meta_text,
                    ft.Container(
                        bgcolor=palette["surface_soft"],
                        border_radius=UI_RADIUS["sm"],
                        padding=UI_SPACING["md"],
                        height=180 if mobile else 220,
                        content=ft.Column([
                            ft.Text("Descripción enriquecida", weight=ft.FontWeight.W_700, color=palette["text"]),
                            description_text,
                        ], scroll=ft.ScrollMode.ADAPTIVE),
                    ),
                    ft.Text("Chat de preparación", weight=ft.FontWeight.W_700, color=palette["text"]),
                    ft.Container(
                        height=300 if mobile else 360,
                        bgcolor=palette["surface_soft"],
                        border_radius=UI_RADIUS["sm"],
                        padding=UI_SPACING["sm"],
                        content=messages_column,
                    ),
                    status_text,
                    ft.Row([
                        message_input,
                        mk_button(
                            "Enviar",
                            on_click=send_message,
                            variant="primary",
                            palette=palette,
                        ),
                    ], spacing=8),
                    ft.Row([
                        mk_button(
                            "Volver al menú",
                            on_click=go_back,
                            variant="secondary",
                            palette=palette,
                        )
                    ]),
                ], spacing=10),
            ),
        ], expand=True, spacing=12)
    )

# helper to show main(job search) UI after auth
async def show_main_ui(page: ft.Page):
    page.controls.clear()
    page.overlay.clear()
    page.title = "JobFriends"
    page.scroll = ft.ScrollMode.ADAPTIVE 
    mobile = is_mobile(page)
    tablet = is_tablet(page)
    page.padding = get_page_padding(page)
    if page.theme_mode == ft.ThemeMode.SYSTEM:
        page.theme_mode = ft.ThemeMode.LIGHT
    palette = get_palette(page)
    page.bgcolor = palette["bg"]
    set_authenticated_bottom_appbar(page, palette, active_section="search")

    async def toggle_theme(e):
        page.theme_mode = ft.ThemeMode.LIGHT if is_dark_mode(page) else ft.ThemeMode.DARK
        await show_main_ui(page)

    page.on_resized = make_responsive_resize_handler(page, lambda: show_main_ui(page))

    # personalize greeting with nombre/apellido if available
    if current_user:
        info = user_details.get(current_user, {})
        nombre = info.get("nombre", "")
        apellido = info.get("apellido", "")
        welcome_text = "Bienvenido"
        if nombre or apellido:
            welcome_text += f" {nombre} {apellido}".strip()
    else:
        welcome_text = "Bienvenido"

    # --- Campos de entrada ---
    kw_input = mk_textfield(
        "Palabras clave",
        palette=palette,
        value="",
        expand=True,
    )
    city_input = mk_textfield(
        "Ciudad",
        palette=palette,
        value="",
        expand=True,
    )
    country_input = mk_textfield(
        "País (co, mx, es)",
        palette=palette,
        value="co",
        width=220,
    )
    results_table = ft.DataTable(
        data_row_max_height=float("inf"),  # permitir filas de altura variable
        column_spacing=12 if mobile else 20,
        horizontal_margin=8 if mobile else 16,
        columns=[
            ft.DataColumn(ft.Text("Vacante")),
            ft.DataColumn(ft.Text("Empresa / Sitio")),
            ft.DataColumn(ft.Text("Ubicación")),
            ft.DataColumn(ft.Text("Acciones")),
        ],
        rows=[]
    )
    results_mobile_list = ft.Column(spacing=UI_SPACING["sm"], expand=True, scroll=ft.ScrollMode.ADAPTIVE)
    results_table_container = ft.Container(
        visible=True,
        expand=True,
        content=ft.Row([
            results_table,
        ], scroll=ft.ScrollMode.ALWAYS),
    )
    results_empty_desktop = ft.Container(
        visible=False,
        expand=True,
        alignment=ft.Alignment.CENTER,
        content=mk_empty_state(
            page,
            palette,
            "No encontramos vacantes para esta búsqueda.",
            "Prueba con otra ciudad o ajusta las palabras clave.",
        ),
    )
    
    # Crear botón de búsqueda como variable para poder deshabilitarlo
    search_button = mk_button("Realizar Búsqueda", variant="primary", palette=palette)
    # Deshabilitar el botón inicialmente
    search_button.disabled = True

    loading_overlay = mk_loading_overlay(
        page,
        palette,
        "Buscando vacantes",
        "Estamos consultando resultados y enriqueciendo la información.",
    )

    # Funciones definidas previamente dentro de main
    async def open_apply_link(e):
        link = e.control.data
        if link and link != "#":
            await ft.UrlLauncher().launch_url(link)

    # Función para validar y habilitar/deshabilitar el botón de búsqueda
    def validate_search_fields(e=None):
        """Valida que los campos obligatorios tengan contenido"""
        kw_filled = kw_input.value.strip() != ""
        city_filled = city_input.value.strip() != ""
        search_button.disabled = not (kw_filled and city_filled)
        page.update()

    # Asignar listeners a los campos de texto
    kw_input.on_change = validate_search_fields
    city_input.on_change = validate_search_fields

    async def show_enriched_description_screen(job_data: dict):
        detail_palette = get_palette(page)
        title = normalize_text(job_data.get("title", "")) or "Detalle de la vacante"
        description_text = normalize_text(job_data.get("enriched_description", "")) or "Sin descripción enriquecida."
        apply_link = normalize_text(job_data.get("apply_link", ""))

        page.controls.clear()
        page.overlay.clear()
        page.title = "JobFriends - Descripción enriquecida"
        page.bgcolor = detail_palette["bg"]
        page.on_resized = lambda e: page.update()
        set_authenticated_bottom_appbar(page, detail_palette, active_section="search")

        page.add(
            ft.Column([
                mk_surface_card(
                    page,
                    detail_palette,
                    content=ft.Row([
                        ft.Text("Descripción Enriquecida", weight=ft.FontWeight.W_900, size=ui_font_size(page, "title"), color=detail_palette["text"]),
                        make_theme_toggle(page, toggle_theme),
                    ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
                ),
                mk_surface_card(
                    page,
                    detail_palette,
                    content=ft.Column([
                        ft.Text(title, size=ui_font_size(page, "subtitle"), weight=ft.FontWeight.W_700, color=detail_palette["text"]),
                        ft.Container(
                            height=420 if not mobile else 320,
                            bgcolor=detail_palette["surface_soft"],
                            border_radius=UI_RADIUS["sm"],
                            padding=UI_SPACING["md"],
                            content=ft.Column([
                                ft.Text(description_text, selectable=True, color=detail_palette["text"]),
                            ], scroll=ft.ScrollMode.ADAPTIVE),
                        ),
                        ft.Row([
                            mk_button(
                                "Aplicar",
                                on_click=open_apply_link,
                                data=apply_link,
                                variant="primary",
                                palette=detail_palette,
                            ),
                            mk_button(
                                "Volver a resultados",
                                on_click=lambda e: asyncio.create_task(show_main_ui(page)),
                                variant="secondary",
                                palette=detail_palette,
                            ),
                        ], wrap=mobile, spacing=UI_SPACING["sm"]),
                    ], spacing=UI_SPACING["md"]),
                ),
            ], spacing=UI_SPACING["md"], expand=True)
        )
        page.update()

    async def open_details(e):
        await show_enriched_description_screen(e.control.data or {})

    async def apply_job(e):
        """Handler called cuando el usuario pulsa Aplicar en una fila de resultados."""
        if not current_user:
            # debería haberse autenticado antes
            show_feedback(page, "Debe iniciar sesión para aplicar.", tone="warning")
            return
        info = e.control.data or {}
        sitio = normalize_text(info.get("website", ""))
        titulo = normalize_text(info.get("title", ""))
        enlace = normalize_text(info.get("apply_link", ""))
        descripcion_enriquecida = normalize_text(info.get("enriched_description", ""))
        descripcion_payload = build_application_description_payload(info, descripcion_enriquecida)
        estado = "Aplicado"

        try:
            payload = {
                "applicant_email": current_user,
                "website": sitio,
                "vaccancy": titulo,
                "status": estado,
                "application_link": enlace,
                "Description": descripcion_payload,
            }
            existing = supabase.table("Applications").select("id").eq("applicant_email", current_user).eq("application_link", enlace).limit(1).execute()
            if existing.data:
                application_id = existing.data[0].get("id")
                supabase.table("Applications").update(payload).eq("id", application_id).execute()
                success_text = "La aplicación existente se actualizó correctamente."
            else:
                supabase.table("Applications").insert(payload).execute()
                success_text = "La aplicación se ha registrado."

            show_feedback(page, success_text, tone="success")
            # abrir la URL en el navegador predeterminado
            if enlace:
                await ft.UrlLauncher().launch_url(enlace)
        except Exception:
            show_feedback(page, "No se pudo registrar la aplicación.", tone="error")

    def perform_serp_search(keywords, city, country_name):
        return search_google_jobs(keywords, city=city, country_name=country_name)

    async def on_search_click(e):
        # 1. Mostrar estado de carga
        loading_overlay.visible = True
        results_empty_desktop.visible = False
        results_table_container.visible = True
        page.update()

        results_table.rows.clear()
        results_mobile_list.controls.clear()

        # 2. Obtener datos de los inputs definidos en tu código
        query = kw_input.value
        ciudad = city_input.value
        pais = "Colombia"
        
        # 3. Llamar a la API
        jobs = await asyncio.to_thread(perform_serp_search, query, ciudad, pais)
        
        if not jobs:
            if mobile:
                results_mobile_list.controls.append(
                    mk_surface_card(
                        page,
                        palette,
                        mk_empty_state(
                            page,
                            palette,
                            "No encontramos vacantes para esta búsqueda.",
                            "Prueba con otra ciudad o ajusta las palabras clave.",
                        ),
                    )
                )
            else:
                results_empty_desktop.visible = True
                results_table_container.visible = False

        else:
            results_empty_desktop.visible = False
            results_table_container.visible = True
            semaphore = asyncio.Semaphore(4)

            async def enrich_job(job: dict):
                apply_options = job.get("apply_options") or job.get("applly_options") or []
                apply_link = apply_options[0].get("link", "#") if apply_options else "#"
                async with semaphore:
                    description = await asyncio.to_thread(build_enriched_description, job, apply_link)
                return apply_link, description

            enriched_results = await asyncio.gather(
                *(enrich_job(job) for job in jobs),
                return_exceptions=True,
            )

            for job, enriched in zip(jobs, enriched_results):
                company = job.get("company_name", "N/A")
                location = job.get("location", "N/A")
                titulo_cargo = normalize_text(job.get("title", "Sin descripción"))

                if isinstance(enriched, Exception):
                    apply_link = "#"
                    descripcion_completa = titulo_cargo
                else:
                    apply_link, descripcion_completa = enriched

                job_payload = {
                    "title": titulo_cargo,
                    "company": normalize_text(company),
                    "website": normalize_text(company),
                    "location": normalize_text(location),
                    "apply_link": apply_link,
                    "share_link": normalize_text(job.get("share_link", "")),
                    "enriched_description": descripcion_completa,
                }
            
                if mobile:
                    results_mobile_list.controls.append(
                        mk_surface_card(
                            page,
                            palette,
                            ft.Column(
                                [
                                    ft.Text(titulo_cargo, color=palette["text"], weight=ft.FontWeight.W_700, size=ui_font_size(page, "subtitle")),
                                    ft.Text(f"Empresa: {company}", color=palette["muted"], size=ui_font_size(page, "label")),
                                    ft.Text(f"Ubicación: {location}", color=palette["muted"], size=ui_font_size(page, "label")),
                                    ft.Row(
                                        [
                                            mk_button("Aplicar", on_click=apply_job, data=job_payload, variant="primary", palette=palette),
                                            mk_button("Ver detalles", on_click=open_details, data=job_payload, variant="secondary", palette=palette),
                                        ],
                                        wrap=True,
                                        spacing=UI_SPACING["sm"],
                                    ),
                                ],
                                spacing=UI_SPACING["sm"],
                            ),
                        )
                    )
                else:
                    results_table.rows.append(
                        ft.DataRow(
                            cells=[
                                ft.DataCell(ft.Text(titulo_cargo)),
                                ft.DataCell(ft.Text(company)),
                                ft.DataCell(ft.Text(location)),
                                ft.DataCell(
                                    ft.Row([
                                        mk_button(
                                            "Aplicar",
                                            on_click=apply_job,
                                            data=job_payload,
                                            variant="primary",
                                            palette=palette,
                                        ),
                                        mk_button(
                                            "Ver detalles",
                                            on_click=open_details,
                                            data=job_payload,
                                            variant="secondary",
                                            palette=palette,
                                        ),
                                    ], wrap=True, spacing=UI_SPACING["sm"])
                                ),
                            ]
                        )
                    )
        
        # 5. Ocultar carga y actualizar
        loading_overlay.visible = False
        page.update()

    # Asignar la función al botón de búsqueda existente
    search_button.on_click = on_search_click        

    async def do_logout(e):
        await logout_user(page)

    async def go_menu(e):
        # regresar al menú principal
        await show_menu_ui(page)

    # Asignar el on_click handler al botón de búsqueda
    search_button.on_click = on_search_click

    # Stack para superponer el overlay de carga sobre el contenido
    results_height = 300 if mobile else 380 if tablet else 520

    main_content = ft.Column([
        mk_surface_card(
            page,
            palette,
            variant="hero",
            content=ft.ResponsiveRow([
                ft.Column([
                    mk_section_header(page, palette, "Búsqueda Inteligente", welcome_text),
                ], spacing=UI_SPACING["xs"], col={"xs": 12, "md": 6}),
                ft.Container(
                    col={"xs": 12, "md": 6},
                    content=ft.Row([
                        make_theme_toggle(page, toggle_theme),
                        mk_button(
                            "Volver al menú",
                            on_click=go_menu,
                            variant="secondary",
                            palette=palette,
                        ),
                        mk_button(
                            "Cerrar sesión",
                            on_click=do_logout,
                            variant="danger",
                            palette=palette,
                        ),
                    ], wrap=mobile, spacing=8, alignment=ft.MainAxisAlignment.START if mobile else ft.MainAxisAlignment.END),
                ),
            ], run_spacing=10),
        ),
        mk_surface_card(
            page,
            palette,
            content=ft.Column([
                ft.Text("Filtros", size=ui_font_size(page, "subtitle"), weight=ft.FontWeight.W_700, color=palette["text"]),
                ft.ResponsiveRow([
                    ft.Container(content=kw_input, col={"xs": 12, "md": 6}),
                    ft.Container(content=city_input, col={"xs": 12, "md": 6}),
                    ft.Container(content=country_input, col={"xs": 12, "md": 4}),
                    ft.Container(content=search_button, col={"xs": 12, "md": 8}),
                ], run_spacing=10),
            ], spacing=12),
        ),
        mk_surface_card(
            page,
            palette,
            content=ft.Column([
                ft.Text("Resultados", size=ui_font_size(page, "subtitle"), weight=ft.FontWeight.W_700, color=palette["text"]),
                ft.Container(
                    height=results_height,
                    content=(
                        results_mobile_list
                        if mobile
                        else ft.Column([
                            results_empty_desktop,
                            results_table_container,
                        ], expand=True)
                    ),
                ),
            ], spacing=12),
        )
    ], spacing=12 if mobile else 14)

    page.add(
        ft.Stack([
            main_content,
            loading_overlay,
        ], expand=True)
    )

# menú de funciones principal
async def show_menu_ui(page: ft.Page):
    page.controls.clear()
    page.overlay.clear()
    page.title = "JobFriends - Menú"
    page.scroll = ft.ScrollMode.ADAPTIVE 
    mobile = is_mobile(page)
    tablet = is_tablet(page)
    page.padding = get_page_padding(page)
    if page.theme_mode == ft.ThemeMode.SYSTEM:
        page.theme_mode = ft.ThemeMode.LIGHT
    palette = get_palette(page)
    page.bgcolor = palette["bg"]
    set_authenticated_bottom_appbar(page, palette)

    async def toggle_theme(e):
        page.theme_mode = ft.ThemeMode.LIGHT if is_dark_mode(page) else ft.ThemeMode.DARK
        await show_menu_ui(page)

    page.on_resized = make_responsive_resize_handler(page, lambda: show_menu_ui(page))

    # obtener datos del usuario logueado
    if current_user:
        info = user_details.get(current_user, {})
        nombre = info.get("nombre", "")
        apellido = info.get("apellido", "")
        welcome_name = f"{nombre} {apellido}".strip()
    else:
        welcome_name = "Usuario"

    async def go_job_search(e):
        await show_main_ui(page)

    async def do_logout(e):
        await logout_user(page)

    cv_record = get_applicant_cv(current_user) if current_user else None
    cv_path_input = mk_textfield(
        "Ruta local del CV (.pdf, .doc, .docx)",
        palette=palette,
    )
    cv_status = mk_status_text(page, palette, "Puedes reemplazar tu CV actual o crear uno nuevo con Gemini.")
    cv_info = ft.Column(spacing=8)

    def refresh_cv_info(record: dict | None = None):
        current_record = record if record is not None else cv_record
        cv_info.controls.clear()
        if current_record:
            source_label = "Gemini" if current_record.get("source") == "ai_generated" else "Archivo cargado"
            cv_info.controls.extend([
                ft.Text(f"CV actual: {current_record.get('file_name', 'Sin nombre')}", color=palette["text"], weight=ft.FontWeight.W_700),
                ft.Text(f"Origen: {source_label}", color=palette["muted"]),
                ft.Text(f"Vacantes objetivo: {current_record.get('target_roles') or 'No especificadas'}", color=palette["muted"]),
            ])
            public_url = current_record.get("public_url", "")
            if public_url:
                cv_info.controls.append(ft.TextButton("Abrir CV guardado", url=public_url))
        else:
            cv_info.controls.append(
                mk_empty_state(
                    page,
                    palette,
                    "Todavía no tienes un CV guardado.",
                    "Puedes reemplazar tu CV actual o crear uno nuevo con Gemini cuando quieras.",
                )
            )

    async def save_uploaded_cv():
        nonlocal cv_record
        file_path = (cv_path_input.value or "").strip().strip('"')
        if not file_path:
            cv_status.value = "Ingresa la ruta local del archivo PDF, DOC o DOCX."
            cv_status.color = palette["danger"]
            page.update()
            return

        if not os.path.isfile(file_path):
            cv_status.value = "La ruta indicada no existe o no es un archivo válido."
            cv_status.color = palette["danger"]
            page.update()
            return

        file_name = Path(file_path).name
        if not is_supported_cv_file(file_name):
            cv_status.value = "Solo se admiten archivos PDF, DOC o DOCX."
            cv_status.color = palette["danger"]
            page.update()
            return

        try:
            metadata = await asyncio.to_thread(upload_cv_file, current_user, file_path, file_name)
            cv_record = metadata
            cv_status.value = "CV reemplazado y guardado correctamente."
            cv_status.color = palette["success"]
            refresh_cv_info(cv_record)
        except Exception:
            cv_status.value = "No se pudo guardar el CV en Supabase."
            cv_status.color = palette["danger"]
        page.update()

    async def save_ai_cv(profile_data: dict[str, str]):
        nonlocal cv_record
        try:
            details = user_details.get(current_user or "", {})
            full_name = f"{details.get('nombre', '')} {details.get('apellido', '')}".strip()
            output_format = normalize_ai_cv_output_format(profile_data.get("output_format", "docx"))
            metadata = await asyncio.to_thread(
                create_ai_cv_for_user,
                current_user,
                full_name or current_user or "Usuario",
                profile_data,
                output_format,
            )
            cv_record = metadata
            cv_status.value = f"Nuevo CV {output_format.upper()} generado con Gemini y guardado correctamente."
            cv_status.color = palette["success"]
            refresh_cv_info(cv_record)
        except Exception:
            cv_status.value = "No se pudo generar o guardar el CV con Gemini."
            cv_status.color = palette["danger"]
        page.update()

    refresh_cv_info(cv_record)

    # Obtener empleos aplicados desde Supabase
    applied_jobs = []
    status_count = {status: 0 for status in APPLICATION_STATUS_OPTIONS}
    
    try:
        # filtrar por applicant_email en lugar de Email
        res = supabase.table("Applications").select("*").eq("applicant_email", current_user).execute()
        if res.data:
            applied_jobs = res.data
            for job in applied_jobs:
                status = normalize_text(job.get("status", "")) or "Aplicado"
                if status in status_count:
                    status_count[status] += 1
                else:
                    status_count[status] = 1
    except Exception:
        # Si la tabla no existe o hay error, simplemente continuamos sin datos
        applied_jobs = []

    # Crear tabla de empleos aplicados
    applied_jobs_table = ft.DataTable(
        column_spacing=12 if mobile else 20,
        horizontal_margin=8 if mobile else 16,
        columns=[
            ft.DataColumn(ft.Text("Sitio web")),
            ft.DataColumn(ft.Text("Cargo")),
            ft.DataColumn(ft.Text("Estado")),
            ft.DataColumn(ft.Text("Acciones")),
        ],
        rows=[],
        expand=True
    )
    applied_jobs_mobile_list = ft.Column(spacing=UI_SPACING["sm"], expand=True, scroll=ft.ScrollMode.ADAPTIVE)

    async def open_job_url(e):
        if e.control.data:
            await ft.UrlLauncher().launch_url(e.control.data)

    async def delete_job(e):
        job = e.control.data
        job_id = job.get("id")
        try:
            supabase.table("Applications").delete().eq("id", job_id).execute()
            show_feedback(page, "Empleo eliminado exitosamente", tone="success")
            # Recargar la tabla
            await show_menu_ui(page)
        except Exception:
            show_feedback(page, "Error al eliminar el empleo", tone="error")

    async def open_interview_for_job(job: dict):
        global pending_interview_application_id
        pending_interview_application_id = str(job.get("id", "")) if job.get("id") is not None else None
        await show_interview_prep_ui(page)

    def make_prepare_interview_handler(captured_job: dict):
        async def _handler(e):
            await open_interview_for_job(captured_job)
        return _handler

    def make_status_change_handler(captured_job: dict):
        async def _on_status_change(e):
            new_status = normalize_text(e.data or "") or normalize_text(getattr(e.control, "value", "") or "")
            job_id = captured_job.get("id")

            if not job_id or not new_status:
                return

            try:
                supabase.table("Applications").update({"status": new_status}).eq("id", job_id).execute()
                show_feedback(page, f"Estado actualizado a '{new_status}'.", tone="info")

                if new_status != "En proceso":
                    await show_menu_ui(page)
                    return

                prompt_dialog = ft.AlertDialog(
                    title=ft.Text("¿Deseas prepararte para la entrevista?"),
                    content=ft.Text(
                        f"Cambiaste el estado de '{get_application_title(captured_job)}' a 'En proceso'.\n"
                        "¿Quieres ir ahora a Preparar Entrevista para esta oferta?"
                    ),
                    modal=True,
                )

                async def close_and_refresh(_):
                    prompt_dialog.open = False
                    page.update()
                    await show_menu_ui(page)

                async def accept_and_go(_):
                    prompt_dialog.open = False
                    page.update()
                    updated_job = dict(captured_job)
                    updated_job["status"] = new_status
                    await open_interview_for_job(updated_job)

                prompt_dialog.actions = [
                    ft.TextButton("No ahora", on_click=lambda ev: asyncio.create_task(close_and_refresh(ev))),
                    mk_button(
                        "Sí, preparar entrevista",
                        on_click=lambda ev: asyncio.create_task(accept_and_go(ev)),
                        variant="primary",
                        palette=palette,
                    ),
                ]

                page.overlay.append(prompt_dialog)
                prompt_dialog.open = True
                page.update()
            except Exception:
                show_feedback(page, "No se pudo actualizar el estado.", tone="error")

        return _on_status_change

    # Llenar tabla
    for job in applied_jobs:
        sitio = job.get("website", "")
        cargo = get_application_title(job)
        estado = normalize_text(job.get("status", "")) or "Aplicado"
        enlace = job.get("application_link", "#")

        if estado not in APPLICATION_STATUS_OPTIONS:
            estado = "Aplicado"
            status_count[estado] = status_count.get(estado, 0) + 1

        status_dropdown = ft.Dropdown(
            value=estado,
            dense=True,
            options=[ft.dropdown.Option(s) for s in APPLICATION_STATUS_OPTIONS],
            width=170,
        )
        status_dropdown.on_change = make_status_change_handler(job)

        interview_btn = mk_button(
            "Preparar entrevista",
            on_click=make_prepare_interview_handler(job),
            variant="secondary",
            palette=palette,
        )

        delete_btn = ft.IconButton(icon=ft.Icons.DELETE, on_click=delete_job, data=job)

        if mobile:
            applied_jobs_mobile_list.controls.append(
                mk_surface_card(
                    page,
                    palette,
                    ft.Column(
                        [
                            ft.Text(cargo, color=palette["text"], weight=ft.FontWeight.W_700, size=ui_font_size(page, "subtitle")),
                            ft.Text(f"Sitio: {sitio or 'N/A'}", color=palette["muted"], size=ui_font_size(page, "label")),
                            status_dropdown,
                            ft.Row(
                                [
                                    ft.TextButton("Ver empleo", url=enlace),
                                    interview_btn,
                                    delete_btn,
                                ],
                                wrap=True,
                                spacing=UI_SPACING["sm"],
                            ),
                        ],
                        spacing=UI_SPACING["sm"],
                    ),
                )
            )
        else:
            applied_jobs_table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(ft.Text(sitio)),
                        ft.DataCell(ft.Text(cargo)),
                        ft.DataCell(status_dropdown),
                        ft.DataCell(
                            ft.Row([
                                ft.TextButton("Ver empleo", url=enlace),
                                interview_btn,
                                delete_btn,
                            ], wrap=mobile, spacing=6)
                        ),
                    ]
                )
            )

    # Crear visualización de distribución por estado con contenedores
    status_items = []
    colors_map = APPLICATION_STATUS_COLORS
    
    for status, count in status_count.items():
        color = colors_map.get(status, "#9E9E9E")
        status_items.append(mk_status_chip(page, palette, status, count, color))
    
    if not status_items:
        status_items.append(
            mk_empty_state(
                page,
                palette,
                "Aún no tienes empleos aplicados.",
                "Cuando registres una aplicación, aquí verás el resumen por estado.",
            )
        )
    
    status_chart = ft.Column(status_items, spacing=12)

    # Panel central con gráfico y tabla
    central_panel = ft.Column([
        mk_surface_card(
            page,
            palette,
            variant="hero",
            content=ft.Row([
                ft.Row([
                    ft.Container(
                        width=40,
                        height=40,
                        border_radius=UI_RADIUS["md"],
                        alignment=ft.Alignment.CENTER,
                        bgcolor=palette["accent"],
                        content=ft.Icon(ft.Icons.WORK, color=palette["accent_text"]),
                    ),
                    ft.Column([
                        ft.Text("JobFriends", weight=ft.FontWeight.W_900, size=ui_font_size(page, "subtitle"), color=palette["text"]),
                        ft.Text("Career Hub", size=ui_font_size(page, "label"), color=palette["muted"]),
                    ], spacing=0),
                ], spacing=10),
                make_theme_toggle(page, toggle_theme),
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
        ),
        mk_surface_card(
            page,
            palette,
            variant="secondary",
            content=ft.Column([
                mk_section_header(
                    page,
                    palette,
                    f"Bienvenido, {welcome_name}",
                    "Gestiona tus aplicaciones y monitorea su estado.",
                ),
            ], spacing=4),
        ),
        mk_surface_card(
            page,
            palette,
            content=ft.Column([
                ft.Text("Distribución de empleos por estado", weight=ft.FontWeight.W_700, size=ui_font_size(page, "subtitle"), color=palette["text"]),
                status_chart,
            ], spacing=12),
        ),
        mk_surface_card(
            page,
            palette,
            expand=True,
            content=ft.Column([
                ft.Text("Empleos a los que has aplicado", weight=ft.FontWeight.W_700, size=ui_font_size(page, "subtitle"), color=palette["text"]),
                ft.Container(
                    expand=True,
                    content=(
                        applied_jobs_mobile_list
                        if mobile
                        else ft.Row([
                            applied_jobs_table,
                        ], scroll=ft.ScrollMode.ALWAYS)
                    ),
                ),
            ], spacing=12, expand=True),
        ),
    ], expand=True, spacing=12)

    main_layout = ft.Column([
        central_panel,
    ], expand=True, spacing=12)

    page.add(main_layout)

# authentication screens
async def show_login_ui(page: ft.Page, email_prefill: str = "", pwd_prefill: str = "", auto_submit: bool = False):
    page.controls.clear()
    page.overlay.clear()
    page.bottom_appbar = None
    page.title = "JobFriends - Login"
    mobile = is_mobile(page)
    tablet = is_tablet(page)
    page.padding = get_page_padding(page)
    if page.theme_mode == ft.ThemeMode.SYSTEM:
        page.theme_mode = ft.ThemeMode.LIGHT
    palette = get_palette(page)
    page.bgcolor = palette["bg"]

    async def toggle_theme(e):
        page.theme_mode = ft.ThemeMode.LIGHT if is_dark_mode(page) else ft.ThemeMode.DARK
        await show_login_ui(page, email_prefill, pwd_prefill, auto_submit)

    page.on_resized = make_responsive_resize_handler(
        page,
        lambda: show_login_ui(page, email_prefill, pwd_prefill, auto_submit),
    )

    user_field = mk_textfield(
        "Email",
        palette=palette,
        expand=True,
        value=email_prefill,
    )
    pass_field = mk_textfield(
        "Contraseña",
        palette=palette,
        password=True,
        can_reveal_password=True,
        expand=True,
        value=pwd_prefill,
    )
    remember_chk = ft.Checkbox(label="Recordarme", label_style=ft.TextStyle(color=palette["muted"], size=ui_font_size(page, "label")))
    message = mk_status_text(page, palette, tone="error")

    async def do_login(e):
        nonlocal remember_chk
        global current_user
        email = user_field.value.strip().lower() if user_field.value else ""
        pwd = pass_field.value if pass_field.value else ""
        if not email or not pwd:
            message.value = "Complete todos los campos"
            message.color = palette["danger"]
            page.update()
            return

        try:
            hashed = hash_password(pwd)
            res = supabase.table("Aplicants").select("Password,First_Name,Last_Name,Phone").eq("Email", email).execute()
            
            if not res.data:
                message.value = "Credenciales incorrectas"
                message.color = palette["danger"]
            else:
                row = res.data[0]
                if row.get("Password") == hashed:
                    current_user = email
                    user_details[email] = {
                        "nombre": row.get("First_Name", ""),
                        "apellido": row.get("Last_Name", ""),
                        "telefono": row.get("Phone", "") or "",
                    }
                    # guardar si corresponde
                    if remember_chk.value:
                        save_credentials(email, pwd)
                        
                    message.value = "Login exitoso"
                    message.color = palette["success"]
                    page.update()
                    await show_menu_ui(page)
                else:
                    message.value = "Credenciales incorrectas"
                    message.color = palette["danger"]

        except Exception:
            traceback.print_exc()
            message.value = "Error al verificar credenciales"
            message.color = palette["danger"]
        page.update()

    async def go_register(e):
        await show_register_ui(page)

    page.add(
        ft.Column([
            ft.Row([
                ft.Container(expand=True),
                make_theme_toggle(page, toggle_theme),
            ]),
            mk_centered_shell(
                page,
                palette,
                ft.Column([
                        mk_section_header(
                            page,
                            palette,
                            "Bienvenido",
                            "Introduce tus credenciales para continuar.",
                        ),
                        user_field,
                        pass_field,
                        remember_chk,
                        ft.Row([
                            mk_button(
                                "Login",
                                on_click=do_login,
                                variant="primary",
                                palette=palette,
                            ),
                            ft.TextButton("Regístrate", on_click=go_register),
                        ], alignment=ft.MainAxisAlignment.CENTER),
                        message,
                    ], spacing=14),
                width=None if mobile else 460 if tablet else 500,
            )
        ], expand=True)
    )
    # permitir login al pulsar Enter en el campo de contraseña
    pass_field.on_submit = lambda e: asyncio.create_task(do_login(e))

    # si se piden credenciales automáticas, disparar el login
    if auto_submit and email_prefill and pwd_prefill:
        # pequeña demora para que la UI se renderice
        async def _auto():
            await asyncio.sleep(0.1)
            await do_login(None)
        asyncio.create_task(_auto())

async def show_register_ui(page: ft.Page):
    page.controls.clear()
    page.overlay.clear()
    page.bottom_appbar = None
    page.title = "JobFriends - Registro"
    mobile = is_mobile(page)
    tablet = is_tablet(page)
    page.padding = get_page_padding(page)
    if page.theme_mode == ft.ThemeMode.SYSTEM:
        page.theme_mode = ft.ThemeMode.LIGHT
    palette = get_palette(page)
    page.bgcolor = palette["bg"]

    async def toggle_theme(e):
        page.theme_mode = ft.ThemeMode.LIGHT if is_dark_mode(page) else ft.ThemeMode.DARK
        await show_register_ui(page)

    page.on_resized = make_responsive_resize_handler(page, lambda: show_register_ui(page))

    # nuevos campos de nombre y apellido
    name_field = mk_textfield("Nombre", palette=palette, expand=True)
    surname_field = mk_textfield("Apellido", palette=palette, expand=True)
    user_field = mk_textfield("Email", palette=palette, expand=True)
    pass_field = mk_textfield("Contraseña", palette=palette, password=True, can_reveal_password=True, expand=True)
    phone_field = mk_textfield("Número de teléfono", palette=palette, expand=True)
    cv_path_input = mk_textfield("Ruta local del CV (.pdf, .doc, .docx)", palette=palette, expand=True, bgcolor=palette["surface"])
    message = mk_status_text(page, palette, tone="error")
    cv_message = mk_status_text(page, palette, "Debes cargar un CV o crearlo con IA antes de registrarte.")
    cv_state = {"mode": None, "file_path": None, "file_name": "", "profile_data": None, "output_format": "docx"}

    def set_cv_status(text: str, color: str):
        cv_message.value = text
        cv_message.color = color
        page.update()

    def use_cv_from_path(e=None):
        file_path = (cv_path_input.value or "").strip().strip('"')
        if not file_path:
            set_cv_status("Ingresa la ruta local del archivo PDF, DOC o DOCX.", palette["danger"])
            return
        if not os.path.isfile(file_path):
            set_cv_status("La ruta indicada no existe o no es un archivo válido.", palette["danger"])
            return
        file_name = Path(file_path).name
        if not is_supported_cv_file(file_name):
            set_cv_status("Solo se admiten archivos PDF, DOC o DOCX.", palette["danger"])
            return
        cv_state["mode"] = "upload"
        cv_state["file_path"] = file_path
        cv_state["file_name"] = file_name
        cv_state["profile_data"] = None
        set_cv_status(f"Archivo seleccionado: {file_name}", palette["success"])

    async def handle_ai_cv_creation(profile_data: dict[str, str]):
        cv_state["mode"] = "ai"
        cv_state["file_path"] = None
        cv_state["file_name"] = ""
        cv_state["profile_data"] = profile_data
        cv_state["output_format"] = normalize_ai_cv_output_format(profile_data.get("output_format", "docx"))
        target_roles = profile_data.get("target_roles", "vacantes seleccionadas")
        set_cv_status(
            f"Borrador de CV {cv_state['output_format'].upper()} con IA listo para: {target_roles}",
            palette["success"],
        )

    async def do_register(e):
        email = user_field.value.strip().lower()
        pwd = pass_field.value
        nombre = name_field.value.strip()
        apellido = surname_field.value.strip()
        telefono = phone_field.value.strip() if phone_field.value else ""

        if not email or not pwd or not nombre or not apellido or not telefono:
            message.value = "Complete todos los campos"
            message.color = palette["danger"]
        elif cv_state["mode"] not in {"upload", "ai"}:
            message.value = "Debes cargar un curriculum o crearlo con IA antes de registrarte"
            message.color = palette["danger"]
        else:
            try:
                # verificar existencia
                exists = supabase.table("Aplicants").select("Email").eq("Email", email).execute()
                if exists.data:
                    message.value = "Usuario ya existe"
                    message.color = palette["danger"]
                else:
                    hashed = hash_password(pwd)
                    supabase.table("Aplicants").insert({
                        "Email": email,
                        "Password": hashed,
                        "First_Name": nombre,
                        "Last_Name": apellido,
                        "Phone": telefono,
                    }).execute()

                    try:
                        if cv_state["mode"] == "upload":
                            await asyncio.to_thread(upload_cv_file, email, cv_state["file_path"], cv_state["file_name"])
                        else:
                            full_name = f"{nombre} {apellido}".strip()
                            await asyncio.to_thread(
                                create_ai_cv_for_user,
                                email,
                                full_name,
                                cv_state["profile_data"],
                                cv_state.get("output_format", "docx"),
                            )
                    except Exception:
                        supabase.table("Aplicants").delete().eq("Email", email).execute()
                        message.value = "No se pudo guardar el curriculum. Se canceló el registro."
                        message.color = palette["danger"]
                        page.update()
                        return

                    await show_login_ui(page)
            except Exception:
                message.value = "Error al registrar"
                message.color = palette["danger"]
        page.update()

    async def go_login(e):
        await show_login_ui(page)

    page.add(
        ft.Column([
            ft.Row([
                ft.Container(expand=True),
                make_theme_toggle(page, toggle_theme),
            ]),
            mk_centered_shell(
                page,
                palette,
                ft.Column([
                        mk_section_header(
                            page,
                            palette,
                            "Crear Cuenta",
                            "Completa tus datos para registrarte.",
                        ),
                        name_field,
                        surname_field,
                        user_field,
                        pass_field,
                        phone_field,
                        ft.Container(
                            bgcolor=palette["surface_soft"],
                            border_radius=16,
                            padding=14,
                            content=ft.Column([
                                ft.Text("Curriculum", weight=ft.FontWeight.W_700, color=palette["text"]),
                                ft.Text(
                                    "Carga tu CV en PDF/DOC/DOCX o créalo con IA antes de finalizar el registro.",
                                    color=palette["muted"],
                                    size=ui_font_size(page, "label"),
                                ),
                                ft.ResponsiveRow([
                                    ft.Container(
                                        col={"xs": 12, "md": 6},
                                        content=cv_path_input,
                                    ),
                                    ft.Container(
                                        col={"xs": 12, "md": 6},
                                        content=mk_button(
                                            "Usar archivo de ruta",
                                            on_click=use_cv_from_path,
                                            variant="secondary",
                                            palette=palette,
                                        ),
                                    ),
                                    ft.Container(
                                        col={"xs": 12, "md": 6},
                                        content=mk_button(
                                            "Crear con IA",
                                            on_click=lambda e: open_ai_cv_builder_dialog(
                                                page,
                                                palette,
                                                cv_state.get("profile_data") or {},
                                                handle_ai_cv_creation,
                                                generate_and_store=False,
                                            ),
                                            variant="primary",
                                            palette=palette,
                                        ),
                                    ),
                                ], run_spacing=10),
                                cv_message,
                            ], spacing=10),
                        ),
                        ft.Row([
                            mk_button(
                                "Registrar",
                                on_click=do_register,
                                variant="primary",
                                palette=palette,
                            ),
                            ft.TextButton("Volver al login", on_click=go_login)
                        ], alignment=ft.MainAxisAlignment.CENTER),
                        message,
                    ], spacing=14),
                width=None if mobile else 500 if tablet else 540,
            ),
        ], expand=True)
    )

# Mantenemos async def para que la API no bloquee la interfaz
async def main(page: ft.Page):
    # usar el modo de tema del sistema para claro/oscuro
    page.overlay.clear()
    page.bottom_appbar = None
    page.theme_mode = ft.ThemeMode.LIGHT
    page.scroll = ft.ScrollMode.ADAPTIVE 
    mobile = is_mobile(page)
    tablet = is_tablet(page)
    page.padding = get_page_padding(page)
    palette = get_palette(page)
    page.bgcolor = palette["bg"]
    # intentar cargar credenciales guardadas
    creds = load_credentials()
    if creds:
        await show_login_ui(page, creds.get("email", ""), creds.get("password", ""), auto_submit=True)
        return

    # show welcome screen with two options
    async def go_job_seeker(e):
        await show_login_ui(page)

    async def go_employer(e):
        # Funcionalidad en desarrollo
        pass

    async def toggle_theme(e):
        page.theme_mode = ft.ThemeMode.LIGHT if is_dark_mode(page) else ft.ThemeMode.DARK
        await main(page)

    page.on_resized = make_responsive_resize_handler(page, lambda: main(page))

    welcome_column = ft.Column([
        ft.Row([
            ft.Container(expand=True),
            make_theme_toggle(page, toggle_theme),
        ]),
        mk_centered_shell(
            page,
            palette,
            ft.Column([
                    ft.Container(
                        width=72 if mobile else 90,
                        height=72 if mobile else 90,
                        border_radius=18 if mobile else 24,
                        bgcolor=palette["accent"],
                        alignment=ft.Alignment.CENTER,
                        content=ft.Icon(ft.Icons.WORK, color=palette["accent_text"], size=36 if mobile else 46),
                    ),
                    ft.Text("JobFriends", size=ui_font_size(page, "title_large") + (10 if not mobile else 0), weight=ft.FontWeight.W_900, color=palette["text"]),
                    ft.Text("Gestión profesional de aplicaciones laborales.", size=ui_font_size(page, "body") + 1, color=palette["muted"], text_align=ft.TextAlign.CENTER),
                    ft.Container(height=10),
                    mk_button(
                        "Busco empleo",
                        width=None if mobile else 320,
                        height=56,
                        on_click=go_job_seeker,
                        variant="primary",
                        palette=palette,
                    ),
                    mk_button(
                        "Ofrezco empleo",
                        width=None if mobile else 320,
                        height=56,
                        on_click=go_employer,
                        variant="secondary",
                        palette=palette,
                    ),
                ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=12),
            width=None if mobile else 520 if tablet else 560,
        ),
    ], alignment=ft.MainAxisAlignment.START, horizontal_alignment=ft.CrossAxisAlignment.CENTER, expand=True, spacing=12)
    
    welcome_container = ft.Container(
        content=welcome_column,
        alignment=ft.Alignment.CENTER,
        expand=True
    )
    page.add(welcome_container)


if __name__ == "__main__":
    # La nueva forma de ejecutar la app
    ft.run(main)